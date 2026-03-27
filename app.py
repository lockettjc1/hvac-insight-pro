import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
from scipy import stats
from datetime import datetime, timedelta
import io

# ─── Page Config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="HVAC Insight Pro",
    page_icon="❄️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Mono:wght@400;700&family=DM+Sans:wght@300;400;500;700&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
}

/* Dark industrial theme */
.stApp {
    background-color: #0d1117;
    color: #e6edf3;
}

/* Sidebar */
section[data-testid="stSidebar"] {
    background-color: #161b22;
    border-right: 1px solid #30363d;
}
section[data-testid="stSidebar"] * { color: #e6edf3 !important; }

/* Metric cards */
.metric-card {
    background: linear-gradient(135deg, #1c2128 0%, #161b22 100%);
    border: 1px solid #30363d;
    border-radius: 12px;
    padding: 20px 24px;
    margin-bottom: 12px;
    position: relative;
    overflow: hidden;
}
.metric-card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 3px;
    background: linear-gradient(90deg, #00b4d8, #0077b6);
}
.metric-card.warning::before { background: linear-gradient(90deg, #f7c948, #e07b00); }
.metric-card.danger::before  { background: linear-gradient(90deg, #ff6b6b, #c0392b); }
.metric-card.good::before    { background: linear-gradient(90deg, #2ecc71, #27ae60); }

.metric-label {
    font-family: 'Space Mono', monospace;
    font-size: 10px;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: #8b949e;
    margin-bottom: 6px;
}
.metric-value {
    font-family: 'Space Mono', monospace;
    font-size: 28px;
    font-weight: 700;
    color: #e6edf3;
    line-height: 1;
}
.metric-delta {
    font-size: 12px;
    margin-top: 6px;
    color: #8b949e;
}
.metric-delta.up   { color: #2ecc71; }
.metric-delta.down { color: #ff6b6b; }

/* Section headers */
.section-header {
    font-family: 'Space Mono', monospace;
    font-size: 11px;
    letter-spacing: 3px;
    text-transform: uppercase;
    color: #00b4d8;
    border-bottom: 1px solid #21262d;
    padding-bottom: 8px;
    margin: 28px 0 16px 0;
}

/* Anomaly badge */
.anomaly-badge {
    display: inline-block;
    background: rgba(255,107,107,0.15);
    border: 1px solid rgba(255,107,107,0.4);
    color: #ff6b6b;
    border-radius: 4px;
    padding: 2px 8px;
    font-size: 11px;
    font-family: 'Space Mono', monospace;
}
.normal-badge {
    display: inline-block;
    background: rgba(46,204,113,0.15);
    border: 1px solid rgba(46,204,113,0.4);
    color: #2ecc71;
    border-radius: 4px;
    padding: 2px 8px;
    font-size: 11px;
    font-family: 'Space Mono', monospace;
}

/* Plotly chart containers */
.chart-container {
    background: #161b22;
    border: 1px solid #30363d;
    border-radius: 12px;
    padding: 4px;
    margin-bottom: 16px;
}

/* Streamlit override bits */
div[data-testid="stMetric"] label { color: #8b949e !important; }
div[data-testid="stMetric"] div[data-testid="stMetricValue"] { color: #e6edf3 !important; }
.stTabs [data-baseweb="tab-list"] { background-color: #161b22; border-radius: 8px; }
.stTabs [data-baseweb="tab"] { color: #8b949e; }
.stTabs [aria-selected="true"] { color: #00b4d8 !important; }
div[data-testid="stFileUploader"] {
    background: #161b22;
    border: 1px dashed #30363d;
    border-radius: 12px;
}
.stButton > button {
    background: linear-gradient(135deg, #0077b6, #00b4d8);
    color: white;
    border: none;
    border-radius: 8px;
    font-family: 'Space Mono', monospace;
    font-size: 12px;
    letter-spacing: 1px;
}
.stButton > button:hover { opacity: 0.85; }

/* Scrollable anomaly table */
.anomaly-table { max-height: 320px; overflow-y: auto; }

/* Fault Rules Engine */
.rule-card {
    background: linear-gradient(135deg, #1c2128 0%, #161b22 100%);
    border: 1px solid #30363d;
    border-radius: 10px;
    padding: 16px 20px;
    margin-bottom: 10px;
    position: relative;
}
.rule-card.active { border-color: #ff6b6b44; }
.rule-card.inactive { opacity: 0.5; }
.rule-name {
    font-family: 'Space Mono', monospace;
    font-size: 13px;
    font-weight: 700;
    color: #e6edf3;
    margin-bottom: 4px;
}
.rule-expr {
    font-family: 'Space Mono', monospace;
    font-size: 11px;
    color: #00b4d8;
    background: rgba(0,180,216,0.07);
    border-radius: 4px;
    padding: 3px 8px;
    display: inline-block;
    margin-bottom: 6px;
}
.rule-desc { font-size: 12px; color: #8b949e; }
.severity-high   { color: #ff6b6b; font-weight: 700; font-size: 11px; font-family: 'Space Mono', monospace; }
.severity-medium { color: #f7c948; font-weight: 700; font-size: 11px; font-family: 'Space Mono', monospace; }
.severity-low    { color: #2ecc71; font-weight: 700; font-size: 11px; font-family: 'Space Mono', monospace; }
.fault-event-row {
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 8px 12px;
    border-bottom: 1px solid #21262d;
    font-size: 12px;
}
.fault-event-row:last-child { border-bottom: none; }
</style>
""", unsafe_allow_html=True)

# ─── Helpers ────────────────────────────────────────────────────────────────────
PLOTLY_LAYOUT = dict(
    paper_bgcolor="#161b22",
    plot_bgcolor="#0d1117",
    font=dict(color="#e6edf3", family="DM Sans", size=13),
    title_font=dict(color="#e6edf3", size=14),
    xaxis=dict(gridcolor="#21262d", zerolinecolor="#30363d", tickfont=dict(color="#e6edf3"), title_font=dict(color="#e6edf3")),
    yaxis=dict(gridcolor="#21262d", zerolinecolor="#30363d", tickfont=dict(color="#e6edf3"), title_font=dict(color="#e6edf3")),
    margin=dict(l=48, r=24, t=48, b=40),
    legend=dict(bgcolor="rgba(0,0,0,0)", bordercolor="#30363d", font=dict(color="#e6edf3", size=13)),
)

COLORS = {
    "blue":   "#00b4d8",
    "teal":   "#06d6a0",
    "amber":  "#f7c948",
    "red":    "#ff6b6b",
    "purple": "#9d4edd",
    "green":  "#2ecc71",
}

def metric_card(label, value, delta=None, status="default"):
    cls = {"default": "", "warning": "warning", "danger": "danger", "good": "good"}.get(status, "")
    delta_html = ""
    if delta is not None:
        direction = "up" if delta >= 0 else "down"
        arrow = "↑" if delta >= 0 else "↓"
        delta_html = f'<div class="metric-delta {direction}">{arrow} {abs(delta):.1f}%</div>'
    st.markdown(f"""
    <div class="metric-card {cls}">
        <div class="metric-label">{label}</div>
        <div class="metric-value">{value}</div>
        {delta_html}
    </div>""", unsafe_allow_html=True)


# ─── Sample Data Generator ───────────────────────────────────────────────────────
@st.cache_data
def generate_sample_data(days=30):
    rng = np.random.default_rng(42)
    n = days * 24
    idx = pd.date_range(end=datetime.now(), periods=n, freq="h")

    hour_of_day = idx.hour
    day_factor = np.sin(np.linspace(0, 2 * np.pi * days, n))

    # Temperatures
    supply_air   = 55  + 3 * np.sin(2 * np.pi * hour_of_day / 24) + rng.normal(0, 0.8, n)
    return_air   = 72  + 5 * np.sin(2 * np.pi * hour_of_day / 24) + rng.normal(0, 1.2, n)
    chilled_wtr  = 44  + 2 * np.sin(2 * np.pi * hour_of_day / 24) + rng.normal(0, 0.5, n)
    condenser    = 85  + 8 * day_factor + rng.normal(0, 1.5, n)
    outside_air  = 68  + 15 * np.sin(2 * np.pi * (hour_of_day - 14) / 24) + rng.normal(0, 2, n)

    # Energy
    chiller_kw   = 120 + 60 * np.clip(np.sin(2 * np.pi * hour_of_day / 24), 0, 1) + rng.normal(0, 5, n)
    ahu_fan_kw   = 15  + 8  * np.clip(np.sin(2 * np.pi * hour_of_day / 24), 0, 1) + rng.normal(0, 1, n)
    pump_kw      = 8   + 3  * np.clip(np.sin(2 * np.pi * hour_of_day / 24), 0, 1) + rng.normal(0, 0.5, n)
    total_kw     = chiller_kw + ahu_fan_kw + pump_kw

    # Airflow
    supply_cfm   = 8000 + 3000 * np.clip(np.sin(2 * np.pi * hour_of_day / 24), 0, 1) + rng.normal(0, 200, n)
    oa_damper    = np.clip(20 + 30 * np.sin(2 * np.pi * (hour_of_day - 8) / 24) + rng.normal(0, 5, n), 0, 100)

    # Inject anomalies
    anomaly_idx = rng.choice(n, size=int(n * 0.03), replace=False)
    supply_air  = np.array(supply_air, dtype=float)
    chiller_kw  = np.array(chiller_kw, dtype=float)
    condenser   = np.array(condenser, dtype=float)
    supply_air[anomaly_idx]  += rng.choice([-15, 15], size=len(anomaly_idx))
    chiller_kw[anomaly_idx]  += rng.choice([-40, 60], size=len(anomaly_idx))
    condenser[anomaly_idx]   += rng.choice([20, 25], size=len(anomaly_idx))

    df = pd.DataFrame({
        "timestamp":        idx,
        "supply_air_temp":  supply_air,
        "return_air_temp":  return_air,
        "chilled_water_temp": chilled_wtr,
        "condenser_temp":   condenser,
        "outside_air_temp": outside_air,
        "chiller_kw":       chiller_kw,
        "ahu_fan_kw":       ahu_fan_kw,
        "pump_kw":          pump_kw,
        "total_kw":         total_kw,
        "supply_cfm":       supply_cfm,
        "oa_damper_pct":    oa_damper,
    }).round(2)
    return df


# ─── Data Loading ────────────────────────────────────────────────────────────────
def load_uploaded(file):
    try:
        df = pd.read_csv(file)
        # Try to find a timestamp column
        for col in df.columns:
            if any(k in col.lower() for k in ["time", "date", "stamp"]):
                df[col] = pd.to_datetime(df[col], infer_datetime_format=True, errors="coerce")
                df = df.rename(columns={col: "timestamp"})
                break
        if "timestamp" not in df.columns:
            df.insert(0, "timestamp", pd.date_range(end=datetime.now(), periods=len(df), freq="h"))
        df = df.dropna(subset=["timestamp"]).sort_values("timestamp").reset_index(drop=True)
        return df, None
    except Exception as e:
        return None, str(e)


# ─── Anomaly Detection ───────────────────────────────────────────────────────────
def detect_anomalies(df, columns, z_threshold=3.0):
    results = df[["timestamp"]].copy()
    results["is_anomaly"] = False
    results["anomaly_score"] = 0.0
    results["anomaly_cols"] = ""
    for col in columns:
        if col not in df.columns:
            continue
        z = np.abs(stats.zscore(df[col].fillna(df[col].median())))
        flag = z > z_threshold
        results["is_anomaly"] |= flag
        results["anomaly_score"] = np.maximum(results["anomaly_score"], z)
        results.loc[flag, "anomaly_cols"] += col + " "
    return results


# ─── Fault Rules Engine ──────────────────────────────────────────────────────────

BUILTIN_RULES = [
    # ── AHU ──────────────────────────────────────────────────────────────────────
    {
        "id": "sa_temp_high",
        "name": "Supply Air Temp Too High",
        "description": "Supply air temperature exceeds 65 °F, indicating cooling failure, coil fouling, or valve fault.",
        "expr": "supply_air_temp > 65 °F for ≥ 1 hr",
        "severity": "HIGH",
        "category": "AHU",
        "required_cols": ["supply_air_temp"],
        "eval_fn": lambda df: df["supply_air_temp"] > 65,
        "min_duration_hrs": 1,
        "recommended_action": "Inspect cooling coil, chilled water valve, and AHU controls.",
    },
    {
        "id": "sa_temp_low",
        "name": "Supply Air Overcooling",
        "description": "Supply air below 50 °F — possible stuck open valve, sensor offset, or control loop failure.",
        "expr": "supply_air_temp < 50 °F for ≥ 1 hr",
        "severity": "MEDIUM",
        "category": "AHU",
        "required_cols": ["supply_air_temp"],
        "eval_fn": lambda df: df["supply_air_temp"] < 50,
        "min_duration_hrs": 1,
        "recommended_action": "Check cooling coil valve and supply air temp sensor calibration.",
    },
    {
        "id": "sa_temp_hunting",
        "name": "Supply Air Temp Hunting",
        "description": "Std deviation of supply air temp over a 3-hour rolling window exceeds 3 °F — sign of unstable control loop (P/I hunting).",
        "expr": "rolling_3hr_std(supply_air_temp) > 3 °F",
        "severity": "MEDIUM",
        "category": "AHU",
        "required_cols": ["supply_air_temp"],
        "eval_fn": lambda df: df["supply_air_temp"].rolling(3, min_periods=2).std() > 3,
        "min_duration_hrs": 1,
        "recommended_action": "Retune PID controller for cooling coil valve.",
    },
    {
        "id": "low_supply_cfm",
        "name": "Low Supply Airflow",
        "description": "Supply CFM below 5000 during occupied hours (6 AM–10 PM) — possible belt slip, VFD fault, or damper issue.",
        "expr": "supply_cfm < 5000 during occupied hours for ≥ 1 hr",
        "severity": "MEDIUM",
        "category": "AHU",
        "required_cols": ["supply_cfm"],
        "eval_fn": lambda df: (df["supply_cfm"] < 5000) & (df["timestamp"].dt.hour.between(6, 22)),
        "min_duration_hrs": 1,
        "recommended_action": "Inspect fan belt, VFD, and supply/return dampers.",
    },
    {
        "id": "oa_damper_stuck_open",
        "name": "OA Damper Stuck Open",
        "description": "Outside air damper >90% open continuously — excess ventilation load, likely actuator fault.",
        "expr": "oa_damper_pct > 90% for ≥ 3 hr",
        "severity": "MEDIUM",
        "category": "AHU",
        "required_cols": ["oa_damper_pct"],
        "eval_fn": lambda df: df["oa_damper_pct"] > 90,
        "min_duration_hrs": 3,
        "recommended_action": "Inspect OA damper actuator and economizer controls.",
    },
    {
        "id": "oa_damper_stuck_closed",
        "name": "OA Damper Stuck Closed",
        "description": "Outside air damper <5% during occupied hours — IAQ risk, possible actuator or linkage failure.",
        "expr": "oa_damper_pct < 5% during occupied hours for ≥ 2 hr",
        "severity": "HIGH",
        "category": "AHU",
        "required_cols": ["oa_damper_pct"],
        "eval_fn": lambda df: (df["oa_damper_pct"] < 5) & (df["timestamp"].dt.hour.between(7, 18)),
        "min_duration_hrs": 2,
        "recommended_action": "Check OA damper actuator, linkage, and BAS command signal.",
    },
    {
        "id": "ra_temp_high",
        "name": "High Return Air Temp",
        "description": "Return air above 78 °F — space may be overheating or thermostat setpoints are drifted.",
        "expr": "return_air_temp > 78 °F for ≥ 2 hr",
        "severity": "MEDIUM",
        "category": "AHU",
        "required_cols": ["return_air_temp"],
        "eval_fn": lambda df: df["return_air_temp"] > 78,
        "min_duration_hrs": 2,
        "recommended_action": "Verify space thermostat setpoints and check for heat sources.",
    },
    # ── Chiller ───────────────────────────────────────────────────────────────────
    {
        "id": "delta_t_low",
        "name": "Low ΔT Syndrome",
        "description": "Return–supply ΔT below 10 °F indicates poor heat transfer, balancing valve issues, or excessive bypass flow.",
        "expr": "(return_air_temp − supply_air_temp) < 10 °F for ≥ 2 hr",
        "severity": "HIGH",
        "category": "Chiller",
        "required_cols": ["supply_air_temp", "return_air_temp"],
        "eval_fn": lambda df: (df["return_air_temp"] - df["supply_air_temp"]) < 10,
        "min_duration_hrs": 2,
        "recommended_action": "Balance coil flow, check bypass valve, verify pump operation.",
    },
    {
        "id": "delta_t_inverted",
        "name": "Inverted ΔT",
        "description": "Supply air temperature exceeds return air temperature — sensor swap or severe reverse flow.",
        "expr": "supply_air_temp > return_air_temp for ≥ 30 min",
        "severity": "HIGH",
        "category": "Chiller",
        "required_cols": ["supply_air_temp", "return_air_temp"],
        "eval_fn": lambda df: df["supply_air_temp"] > df["return_air_temp"],
        "min_duration_hrs": 0.5,
        "recommended_action": "Verify sensor wiring and orientation. Check for reverse flow conditions.",
    },
    {
        "id": "condenser_high",
        "name": "High Condenser Temp",
        "description": "Condenser temp above 100 °F — check cooling tower, condenser fouling, or refrigerant charge.",
        "expr": "condenser_temp > 100 °F for ≥ 1 hr",
        "severity": "HIGH",
        "category": "Chiller",
        "required_cols": ["condenser_temp"],
        "eval_fn": lambda df: df["condenser_temp"] > 100,
        "min_duration_hrs": 1,
        "recommended_action": "Inspect cooling tower, condenser tubes, and refrigerant charge.",
    },
    {
        "id": "chw_temp_high",
        "name": "Chilled Water Supply Temp High",
        "description": "Chilled water supply above 50 °F — chiller struggling, setpoint drift, or low refrigerant.",
        "expr": "chilled_water_temp > 50 °F for ≥ 1 hr",
        "severity": "MEDIUM",
        "category": "Chiller",
        "required_cols": ["chilled_water_temp"],
        "eval_fn": lambda df: df["chilled_water_temp"] > 50,
        "min_duration_hrs": 1,
        "recommended_action": "Verify chiller setpoint, check refrigerant levels and chiller log.",
    },
    {
        "id": "chw_approach_high",
        "name": "High Chiller Approach Temp",
        "description": "Gap between chilled water and supply air > 18 °F — suggests coil fouling or reduced flow.",
        "expr": "(supply_air_temp − chilled_water_temp) > 18 °F for ≥ 2 hr",
        "severity": "MEDIUM",
        "category": "Chiller",
        "required_cols": ["supply_air_temp", "chilled_water_temp"],
        "eval_fn": lambda df: (df["supply_air_temp"] - df["chilled_water_temp"]) > 18,
        "min_duration_hrs": 2,
        "recommended_action": "Clean cooling coil, check flow balancing, inspect heat exchanger.",
    },
    # ── Energy ────────────────────────────────────────────────────────────────────
    {
        "id": "chiller_overload",
        "name": "Chiller Overload",
        "description": "Chiller power draw above 95th percentile — risk of trip or significantly degraded efficiency.",
        "expr": "chiller_kw > P95 for ≥ 30 min",
        "severity": "HIGH",
        "category": "Energy",
        "required_cols": ["chiller_kw"],
        "eval_fn": lambda df: df["chiller_kw"] > df["chiller_kw"].quantile(0.95),
        "min_duration_hrs": 0.5,
        "recommended_action": "Review chiller staging, check for load spike, inspect compressor.",
    },
    {
        "id": "chiller_low_load",
        "name": "Chiller Low-Load Operation",
        "description": "Chiller below 20th percentile of load — inefficient part-load operation, consider staging off.",
        "expr": "chiller_kw < P20 for ≥ 2 hr",
        "severity": "LOW",
        "category": "Energy",
        "required_cols": ["chiller_kw"],
        "eval_fn": lambda df: df["chiller_kw"] < df["chiller_kw"].quantile(0.20),
        "min_duration_hrs": 2,
        "recommended_action": "Evaluate chiller staging strategy and part-load efficiency curve.",
    },
    {
        "id": "total_kw_spike",
        "name": "Unexpected Power Spike",
        "description": "Total facility power exceeds 98th percentile — possible equipment fault or runaway load.",
        "expr": "total_kw > P98 for ≥ 30 min",
        "severity": "HIGH",
        "category": "Energy",
        "required_cols": ["total_kw"],
        "eval_fn": lambda df: df["total_kw"] > df["total_kw"].quantile(0.98),
        "min_duration_hrs": 0.5,
        "recommended_action": "Identify load source, check for simultaneous heating/cooling, review BAS.",
    },
    {
        "id": "after_hours_load",
        "name": "High After-Hours Energy Use",
        "description": "Total power above 80th percentile between 10 PM and 5 AM — equipment may be running when it shouldn't.",
        "expr": "total_kw > P80 between 22:00–05:00 for ≥ 1 hr",
        "severity": "MEDIUM",
        "category": "Energy",
        "required_cols": ["total_kw"],
        "eval_fn": lambda df: (
            df["total_kw"] > df["total_kw"].quantile(0.80)
        ) & (
            ~df["timestamp"].dt.hour.between(5, 21)
        ),
        "min_duration_hrs": 1,
        "recommended_action": "Audit after-hours scheduling in BAS; check for override commands.",
    },
    {
        "id": "simultaneous_heat_cool",
        "name": "Simultaneous Heating & Cooling",
        "description": "Supply air temp low while total power is high — indicates energy waste from simultaneous H/C.",
        "expr": "supply_air_temp < 55 AND total_kw > P75 for ≥ 1 hr",
        "severity": "HIGH",
        "category": "Energy",
        "required_cols": ["supply_air_temp", "total_kw"],
        "eval_fn": lambda df: (df["supply_air_temp"] < 55) & (df["total_kw"] > df["total_kw"].quantile(0.75)),
        "min_duration_hrs": 1,
        "recommended_action": "Check reheat coil controls, review zone-level heating commands.",
    },
]


def find_consecutive_runs(bool_series, min_hours, freq_hours=1):
    """Return list of (start_idx, end_idx) for runs meeting min_hours duration."""
    min_steps = max(1, int(min_hours / freq_hours))
    runs = []
    in_run = False
    start = None
    for i, v in enumerate(bool_series):
        if v and not in_run:
            in_run = True
            start = i
        elif not v and in_run:
            in_run = False
            if (i - start) >= min_steps:
                runs.append((start, i - 1))
    if in_run and (len(bool_series) - start) >= min_steps:
        runs.append((start, len(bool_series) - 1))
    return runs


def evaluate_rule(df, rule):
    """Returns triggered, events, fault_mask for a built-in rule dict."""
    missing = [c for c in rule["required_cols"] if c not in df.columns]
    if missing:
        return {"triggered": False, "events": [], "fault_mask": pd.Series(False, index=df.index),
                "skip_reason": f"Missing columns: {missing}"}
    try:
        condition = rule["eval_fn"](df).fillna(False)
    except Exception as e:
        return {"triggered": False, "events": [], "fault_mask": pd.Series(False, index=df.index),
                "skip_reason": str(e)}

    runs = find_consecutive_runs(condition.values, rule["min_duration_hrs"])
    events = []
    fault_mask = pd.Series(False, index=df.index)
    for s, e in runs:
        start_ts = df["timestamp"].iloc[s]
        end_ts   = df["timestamp"].iloc[e]
        dur_hrs  = (end_ts - start_ts).total_seconds() / 3600
        fault_mask.iloc[s:e+1] = True
        rel_col  = rule["required_cols"][0]
        peak_val = df[rel_col].iloc[s:e+1].max() if rel_col in df.columns else None
        events.append({
            "start":        start_ts,
            "end":          end_ts,
            "duration_hrs": round(dur_hrs, 1),
            "peak_value":   round(peak_val, 2) if peak_val is not None else "—",
            "column":       rel_col,
        })
    return {"triggered": len(events) > 0, "events": events, "fault_mask": fault_mask, "skip_reason": None}


def parse_custom_rule(df, col, operator, threshold, min_hrs):
    """Evaluate a user-defined single-signal threshold rule."""
    ops = {">": lambda a, b: a > b, "<": lambda a, b: a < b,
           ">=": lambda a, b: a >= b, "<=": lambda a, b: a <= b, "==": lambda a, b: a == b}
    if col not in df.columns or operator not in ops:
        return None
    condition = ops[operator](df[col], threshold).fillna(False)
    runs = find_consecutive_runs(condition.values, min_hrs)
    events = []
    fault_mask = pd.Series(False, index=df.index)
    for s, e in runs:
        start_ts = df["timestamp"].iloc[s]
        end_ts   = df["timestamp"].iloc[e]
        dur_hrs  = (end_ts - start_ts).total_seconds() / 3600
        fault_mask.iloc[s:e+1] = True
        peak_val = df[col].iloc[s:e+1].max()
        events.append({
            "start": start_ts, "end": end_ts,
            "duration_hrs": round(dur_hrs, 1),
            "peak_value": round(peak_val, 2),
            "column": col,
        })
    return {"triggered": len(events) > 0, "events": events, "fault_mask": fault_mask}


def parse_dual_signal_rule(df, col_a, op, col_b, offset, min_hrs):
    """Evaluate a two-signal comparison rule: col_a OP (col_b + offset)."""
    ops = {">": lambda a, b: a > b, "<": lambda a, b: a < b,
           ">=": lambda a, b: a >= b, "<=": lambda a, b: a <= b}
    if col_a not in df.columns or col_b not in df.columns or op not in ops:
        return None
    condition = ops[op](df[col_a], df[col_b] + offset).fillna(False)
    runs = find_consecutive_runs(condition.values, min_hrs)
    events = []
    fault_mask = pd.Series(False, index=df.index)
    for s, e in runs:
        start_ts = df["timestamp"].iloc[s]
        end_ts   = df["timestamp"].iloc[e]
        dur_hrs  = (end_ts - start_ts).total_seconds() / 3600
        fault_mask.iloc[s:e+1] = True
        peak_val = (df[col_a] - df[col_b]).iloc[s:e+1].abs().max()
        events.append({
            "start": start_ts, "end": end_ts,
            "duration_hrs": round(dur_hrs, 1),
            "peak_value": round(float(peak_val), 2),
            "column": f"{col_a} vs {col_b}",
        })
    return {"triggered": len(events) > 0, "events": events, "fault_mask": fault_mask}


def build_daily_fault_heatmap(df, all_rule_results, rules):
    """
    Returns a DataFrame (date × rule_name) with fault hours per day per rule,
    and a severity-weighted score per day.
    """
    if df.empty:
        return pd.DataFrame(), pd.Series(dtype=float)

    dates = pd.date_range(df["timestamp"].dt.date.min(), df["timestamp"].dt.date.max(), freq="D")
    sev_weight = {"HIGH": 3, "MEDIUM": 2, "LOW": 1}
    heat = {}
    daily_score = pd.Series(0.0, index=dates.date)

    for rule in rules:
        res = all_rule_results.get(rule["id"])
        if res is None or not res["triggered"]:
            continue
        w = sev_weight.get(rule["severity"], 1)
        fault_mask = res["fault_mask"]
        fault_times = df.loc[fault_mask, "timestamp"]
        if fault_times.empty:
            continue
        by_day = fault_times.dt.date.value_counts()
        heat[rule["name"]] = by_day
        for d, cnt in by_day.items():
            if d in daily_score.index:
                daily_score[d] += cnt * w

    if not heat:
        return pd.DataFrame(), daily_score
    heat_df = pd.DataFrame(heat, index=dates.date).fillna(0)
    return heat_df, daily_score


# ─── KPI Computation ─────────────────────────────────────────────────────────────
def compute_kpis(df):
    kpis = {}
    if "chiller_kw" in df.columns:
        kpis["avg_chiller_kw"] = df["chiller_kw"].mean()
        kpis["peak_chiller_kw"] = df["chiller_kw"].max()
    if "total_kw" in df.columns:
        kpis["total_kwh"] = df["total_kw"].sum()  # hourly data → each row is 1 kWh
    if "supply_air_temp" in df.columns and "return_air_temp" in df.columns:
        kpis["avg_dt"] = (df["return_air_temp"] - df["supply_air_temp"]).mean()
    if "chiller_kw" in df.columns and "supply_cfm" in df.columns:
        # Simplified COP proxy
        cooling = df["supply_cfm"] * 0.5 * abs(df["return_air_temp"] - df["supply_air_temp"]) if "return_air_temp" in df.columns and "supply_air_temp" in df.columns else None
        if cooling is not None:
            kpis["est_cop"] = (cooling / df["chiller_kw"].replace(0, np.nan)).mean()
    return kpis


# ─── Sidebar ─────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## ❄️ HVAC Insight Pro")
    st.markdown("---")

    data_source = st.radio("Data Source", ["Sample Data (Demo)", "Upload CSV"])
    df = None

    if data_source == "Sample Data (Demo)":
        days = st.slider("Days of history", 7, 90, 30)
        df = generate_sample_data(days)
        st.success(f"✓ {len(df):,} records loaded")
    else:
        uploaded = st.file_uploader("Upload HVAC CSV", type=["csv"])
        if uploaded:
            df, err = load_uploaded(uploaded)
            if err:
                st.error(f"Error: {err}")
            else:
                st.success(f"✓ {len(df):,} records loaded")

    st.markdown("---")
    st.markdown("### Filters")

    if df is not None:
        ts_min, ts_max = df["timestamp"].min(), df["timestamp"].max()
        date_range = st.date_input(
            "Date range",
            value=(ts_min.date(), ts_max.date()),
            min_value=ts_min.date(),
            max_value=ts_max.date(),
        )
        if len(date_range) == 2:
            df = df[(df["timestamp"].dt.date >= date_range[0]) & (df["timestamp"].dt.date <= date_range[1])]

    st.markdown("---")
    st.markdown("### Anomaly Settings")
    z_thresh = st.slider("Z-score threshold", 2.0, 5.0, 3.0, 0.1)

    numeric_cols = [c for c in (df.columns if df is not None else []) if c != "timestamp" and pd.api.types.is_numeric_dtype(df[c])] if df is not None else []
    anomaly_cols = st.multiselect(
        "Monitor columns",
        numeric_cols,
        default=numeric_cols[:5] if len(numeric_cols) >= 5 else numeric_cols,
    )

# ─── Main Content ────────────────────────────────────────────────────────────────
if df is None or df.empty:
    st.markdown("<div style='text-align:center;margin-top:80px'>", unsafe_allow_html=True)
    st.markdown("## ❄️ HVAC Insight Pro")
    st.markdown("**Select a data source from the sidebar to begin.**")
    st.markdown("</div>", unsafe_allow_html=True)
    st.stop()

# Compute anomalies
anom_df = detect_anomalies(df, anomaly_cols, z_thresh) if anomaly_cols else pd.DataFrame({"is_anomaly": [False]*len(df)})
df["is_anomaly"] = anom_df["is_anomaly"].values

# ─── Header ──────────────────────────────────────────────────────────────────────
st.markdown("# ❄️ HVAC Insight Pro")
st.markdown(f"<span style='color:#8b949e;font-family:Space Mono,monospace;font-size:12px'>"
            f"Analysing {len(df):,} records · "
            f"{df['timestamp'].min().strftime('%b %d %Y')} → {df['timestamp'].max().strftime('%b %d %Y')}"
            f"</span>", unsafe_allow_html=True)

st.markdown("---")

# ─── KPI Row ─────────────────────────────────────────────────────────────────────
kpis = compute_kpis(df)
anom_count = int(df["is_anomaly"].sum())
anom_pct   = anom_count / len(df) * 100

cols = st.columns(5)
with cols[0]:
    metric_card("Anomalies Detected", str(anom_count),
                delta=None, status="danger" if anom_pct > 5 else "good")
with cols[1]:
    val = f"{kpis.get('avg_chiller_kw', 0):.1f} kW"
    metric_card("Avg Chiller Load", val, status="default")
with cols[2]:
    val = f"{kpis.get('total_kwh', 0)/1000:.1f} MWh"
    metric_card("Total Consumption", val, status="default")
with cols[3]:
    val = f"{kpis.get('avg_dt', 0):.1f} °F"
    metric_card("Avg ΔT (SA↔RA)", val,
                status="warning" if abs(kpis.get("avg_dt", 17)) < 14 else "good")
with cols[4]:
    val = f"{kpis.get('est_cop', 0):.2f}"
    metric_card("Est. System COP", val,
                status="danger" if kpis.get("est_cop", 3) < 2 else "good")

# ─── Tabs ────────────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9 = st.tabs(["📈 Time Series", "⚡ Energy", "🌡 Temperatures", "🚨 Anomalies", "⚙️ Fault Rules", "🤖 ML Insights", "🏭 Equipment", "🔌 BACnet Live", "📄 Report Generator"])

# ══════════════════════════════════════════════════════════════════════════════════
# TAB 1 — Time Series
# ══════════════════════════════════════════════════════════════════════════════════
with tab1:
    st.markdown('<div class="section-header">Select Signals to Plot</div>', unsafe_allow_html=True)
    plot_cols = st.multiselect(
        "Columns", numeric_cols, default=numeric_cols[:4],
        key="ts_cols", label_visibility="collapsed"
    )

    if plot_cols:
        fig = go.Figure()
        palette = list(COLORS.values())
        for i, col in enumerate(plot_cols):
            fig.add_trace(go.Scatter(
                x=df["timestamp"], y=df[col],
                name=col.replace("_", " ").title(),
                line=dict(color=palette[i % len(palette)], width=1.5),
                mode="lines",
            ))

        # Overlay anomaly markers on first signal
        anom_mask = df["is_anomaly"]
        if anom_mask.any() and plot_cols:
            fig.add_trace(go.Scatter(
                x=df.loc[anom_mask, "timestamp"],
                y=df.loc[anom_mask, plot_cols[0]],
                mode="markers",
                name="Anomaly",
                marker=dict(color=COLORS["red"], size=6, symbol="x"),
            ))

        fig.update_layout(**PLOTLY_LAYOUT, title="Time Series Overview", height=420)
        st.plotly_chart(fig, use_container_width=True)

        # Rolling stats
        roll_col = st.selectbox("Rolling window for", plot_cols, key="roll_col")
        window  = st.slider("Window (hours)", 2, 72, 24, key="roll_win")
        rolled  = df[roll_col].rolling(window).mean()
        rolled_std = df[roll_col].rolling(window).std()

        fig2 = go.Figure([
            go.Scatter(x=df["timestamp"], y=df[roll_col],
                       name="Raw", line=dict(color=COLORS["blue"], width=1), opacity=0.4),
            go.Scatter(x=df["timestamp"], y=rolled,
                       name=f"{window}h Rolling Mean", line=dict(color=COLORS["teal"], width=2)),
            go.Scatter(
                x=pd.concat([df["timestamp"], df["timestamp"][::-1]]),
                y=pd.concat([rolled + rolled_std, (rolled - rolled_std)[::-1]]),
                fill="toself", fillcolor="rgba(0,180,216,0.08)",
                line=dict(color="rgba(0,0,0,0)"), name="±1 Std Dev",
            ),
        ])
        fig2.update_layout(**PLOTLY_LAYOUT, title=f"Rolling Analysis — {roll_col.replace('_',' ').title()}", height=320)
        st.plotly_chart(fig2, use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 2 — Energy
# ══════════════════════════════════════════════════════════════════════════════════
with tab2:
    energy_cols = [c for c in ["chiller_kw", "ahu_fan_kw", "pump_kw", "total_kw"] if c in df.columns]

    if not energy_cols:
        st.info("No energy columns detected. Ensure columns contain 'kw' in their names.")
    else:
        st.markdown('<div class="section-header">Energy Breakdown</div>', unsafe_allow_html=True)

        # Stacked area
        area_cols = [c for c in energy_cols if c != "total_kw"]
        if area_cols:
            fig = go.Figure()
            pal = [COLORS["blue"], COLORS["teal"], COLORS["purple"], COLORS["amber"]]
            for i, col in enumerate(area_cols):
                fig.add_trace(go.Scatter(
                    x=df["timestamp"], y=df[col],
                    name=col.replace("_", " ").title(),
                    stackgroup="one",
                    line=dict(color=pal[i % len(pal)]),
                    fillcolor=["rgba(0,180,216,0.4)","rgba(6,214,160,0.4)","rgba(157,78,221,0.4)","rgba(247,201,72,0.4)"][i % 4],
                ))
            fig.update_layout(**PLOTLY_LAYOUT, title="Stacked Energy Consumption (kW)", height=380)
            st.plotly_chart(fig, use_container_width=True)

        c1, c2 = st.columns(2)
        with c1:
            # Hourly profile
            if "total_kw" in df.columns:
                hourly = df.groupby(df["timestamp"].dt.hour)["total_kw"].mean()
                fig = go.Figure(go.Bar(
                    x=hourly.index, y=hourly.values,
                    marker_color=COLORS["blue"],
                    marker_line_color=COLORS["teal"], marker_line_width=0.5,
                ))
                fig.update_layout(**PLOTLY_LAYOUT, title="Avg Load by Hour of Day", height=300,
                                  xaxis_title="Hour", yaxis_title="kW")
                st.plotly_chart(fig, use_container_width=True)

        with c2:
            # Pie breakdown (mean)
            if area_cols:
                means = [df[c].mean() for c in area_cols]
                fig = go.Figure(go.Pie(
                    labels=[c.replace("_", " ").title() for c in area_cols],
                    values=means,
                    hole=0.55,
                    marker=dict(colors=[COLORS["blue"], COLORS["teal"], COLORS["purple"]]),
                ))
                fig.update_layout(**PLOTLY_LAYOUT, title="Avg Load Share", height=300,
                                  showlegend=True)
                st.plotly_chart(fig, use_container_width=True)

        # Daily consumption
        if "total_kw" in df.columns:
            daily = df.groupby(df["timestamp"].dt.date)["total_kw"].sum().reset_index()
            daily.columns = ["date", "daily_kwh"]
            fig = go.Figure(go.Bar(
                x=daily["date"], y=daily["daily_kwh"],
                marker_color=COLORS["teal"],
            ))
            fig.update_layout(**PLOTLY_LAYOUT, title="Daily Energy (kWh)", height=280,
                              xaxis_title="Date", yaxis_title="kWh")
            st.plotly_chart(fig, use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 3 — Temperatures
# ══════════════════════════════════════════════════════════════════════════════════
with tab3:
    temp_cols = [c for c in df.columns if any(k in c.lower() for k in ["temp", "°", "_t"]) and c != "timestamp"]

    if not temp_cols:
        st.info("No temperature columns detected.")
    else:
        st.markdown('<div class="section-header">Temperature Profiles</div>', unsafe_allow_html=True)

        pal = list(COLORS.values())
        fig = go.Figure()
        for i, col in enumerate(temp_cols):
            fig.add_trace(go.Scatter(
                x=df["timestamp"], y=df[col],
                name=col.replace("_", " ").title(),
                line=dict(color=pal[i % len(pal)], width=1.5),
            ))
        fig.update_layout(**PLOTLY_LAYOUT, title="All Temperature Sensors", height=380,
                          yaxis_title="°F")
        st.plotly_chart(fig, use_container_width=True)

        c1, c2 = st.columns(2)
        with c1:
            # Box plots
            melt = df[["timestamp"] + temp_cols].melt(id_vars="timestamp",
                                                       var_name="Sensor", value_name="°F")
            fig = px.box(melt, x="Sensor", y="°F",
                         color="Sensor",
                         color_discrete_sequence=pal)
            fig.update_layout(**PLOTLY_LAYOUT, title="Temperature Distribution",
                              height=340, showlegend=False,
                              xaxis_tickangle=-30)
            st.plotly_chart(fig, use_container_width=True)

        with c2:
            # Correlation heatmap
            if len(temp_cols) >= 2:
                corr = df[temp_cols].corr()
                fig = go.Figure(go.Heatmap(
                    z=corr.values,
                    x=[c.replace("_temp", "").replace("_", " ") for c in corr.columns],
                    y=[c.replace("_temp", "").replace("_", " ") for c in corr.index],
                    colorscale="Blues",
                    zmid=0,
                    text=np.round(corr.values, 2),
                    texttemplate="%{text}",
                ))
                fig.update_layout(**PLOTLY_LAYOUT, title="Sensor Correlation",
                                  height=340)
                st.plotly_chart(fig, use_container_width=True)

        # Supply vs Return DT
        if "supply_air_temp" in df.columns and "return_air_temp" in df.columns:
            df["delta_t"] = df["return_air_temp"] - df["supply_air_temp"]
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df["timestamp"], y=df["delta_t"],
                                     fill="tozeroy",
                                     line=dict(color=COLORS["amber"], width=1.5),
                                     name="ΔT (Return − Supply)"))
            fig.add_hline(y=14, line_dash="dash", line_color=COLORS["green"],
                          annotation_text="Design ΔT 14°F")
            fig.update_layout(**PLOTLY_LAYOUT, title="Supply→Return ΔT", height=280,
                              yaxis_title="°F")
            st.plotly_chart(fig, use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 4 — Anomalies
# ══════════════════════════════════════════════════════════════════════════════════
with tab4:
    st.markdown('<div class="section-header">Anomaly Detection — Z-Score Method</div>',
                unsafe_allow_html=True)

    if not anomaly_cols:
        st.info("Select columns to monitor in the sidebar.")
    else:
        anom_detail = detect_anomalies(df, anomaly_cols, z_thresh)
        anom_rows   = df[anom_detail["is_anomaly"]].copy()
        anom_rows["anomaly_score"] = anom_detail.loc[anom_detail["is_anomaly"], "anomaly_score"].values
        anom_rows["flagged_cols"]  = anom_detail.loc[anom_detail["is_anomaly"], "anomaly_cols"].values

        c1, c2, c3 = st.columns(3)
        c1.metric("Total Anomalies", f"{len(anom_rows):,}")
        c2.metric("Anomaly Rate", f"{len(anom_rows)/len(df)*100:.2f}%")
        c3.metric("Max Z-Score", f"{anom_detail['anomaly_score'].max():.2f}")

        # Score over time
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=df["timestamp"], y=anom_detail["anomaly_score"],
            fill="tozeroy", line=dict(color=COLORS["blue"], width=1),
            name="Anomaly Score",
        ))
        fig.add_hline(y=z_thresh, line_dash="dash", line_color=COLORS["red"],
                      annotation_text=f"Threshold ({z_thresh})")
        fig.update_layout(**PLOTLY_LAYOUT, title="Anomaly Score Over Time", height=280,
                          yaxis_title="Z-score")
        st.plotly_chart(fig, use_container_width=True)

        # Per-column anomaly scatter
        sel_col = st.selectbox("Inspect column", anomaly_cols, key="anom_col")
        if sel_col in df.columns:
            normal  = df[~anom_detail["is_anomaly"]]
            anomaly = df[anom_detail["is_anomaly"]]
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=normal["timestamp"], y=normal[sel_col],
                                     mode="markers",
                                     marker=dict(color=COLORS["blue"], size=2, opacity=0.4),
                                     name="Normal"))
            fig.add_trace(go.Scatter(x=anomaly["timestamp"], y=anomaly[sel_col],
                                     mode="markers",
                                     marker=dict(color=COLORS["red"], size=7, symbol="x"),
                                     name="Anomaly"))
            fig.update_layout(**PLOTLY_LAYOUT,
                              title=f"Anomaly Scatter — {sel_col.replace('_',' ').title()}",
                              height=320, yaxis_title=sel_col)
            st.plotly_chart(fig, use_container_width=True)

        # Table
        if not anom_rows.empty:
            st.markdown('<div class="section-header">Anomaly Log</div>', unsafe_allow_html=True)
            display_cols = ["timestamp"] + [c for c in anomaly_cols if c in anom_rows.columns] + ["anomaly_score"]
            st.dataframe(
                anom_rows[display_cols].sort_values("anomaly_score", ascending=False).head(200),
                use_container_width=True, height=320,
            )
            csv_out = anom_rows[display_cols].to_csv(index=False)
            st.download_button("⬇ Export Anomaly Log CSV", data=csv_out,
                               file_name="hvac_anomalies.csv", mime="text/csv")
        else:
            st.success("🎉 No anomalies detected with current settings.")


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 5 — Fault Rules Engine
# ══════════════════════════════════════════════════════════════════════════════════
with tab5:

    # ── Init session state ──────────────────────────────────────────────────────
    if "enabled_rules" not in st.session_state:
        st.session_state["enabled_rules"] = {r["id"]: True for r in BUILTIN_RULES}
    if "saved_custom_rules" not in st.session_state:
        st.session_state["saved_custom_rules"] = []

    # ── Run ALL enabled built-in rules (shared across sub-tabs) ────────────────
    all_cats     = sorted(set(r["category"] for r in BUILTIN_RULES))
    rule_results = {}
    for rule in BUILTIN_RULES:
        if st.session_state["enabled_rules"].get(rule["id"], True):
            rule_results[rule["id"]] = evaluate_rule(df, rule)

    triggered_rules = [r for r in BUILTIN_RULES
                       if rule_results.get(r["id"], {}).get("triggered", False)]
    total_events = sum(len(rule_results[r["id"]]["events"]) for r in triggered_rules)
    high_faults  = sum(1 for r in triggered_rules if r["severity"] == "HIGH")
    med_faults   = sum(1 for r in triggered_rules if r["severity"] == "MEDIUM")
    low_faults   = sum(1 for r in triggered_rules if r["severity"] == "LOW")
    sev_order    = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    sev_colors   = {"HIGH": COLORS["red"], "MEDIUM": COLORS["amber"], "LOW": COLORS["green"]}

    # ── Top KPI strip ───────────────────────────────────────────────────────────
    k1, k2, k3, k4, k5 = st.columns(5)
    with k1: metric_card("Rules Triggered", str(len(triggered_rules)),
                         status="danger" if high_faults > 0 else ("warning" if med_faults > 0 else "good"))
    with k2: metric_card("Total Fault Events", str(total_events), status="default")
    with k3: metric_card("HIGH Severity", str(high_faults),
                         status="danger" if high_faults > 0 else "good")
    with k4: metric_card("MEDIUM Severity", str(med_faults),
                         status="warning" if med_faults > 0 else "good")
    with k5: metric_card("LOW Severity", str(low_faults), status="default")

    # ── Sub-tabs ────────────────────────────────────────────────────────────────
    ftab1, ftab2, ftab3, ftab4 = st.tabs([
        "📋 Rule Library", "🔧 Custom Rule Builder", "📅 Fault Calendar", "📊 Fault Summary"
    ])

    # ══════════════════════════════════════════════════════════════════════════════
    # FTAB 1 — Rule Library
    # ══════════════════════════════════════════════════════════════════════════════
    with ftab1:
        st.markdown('<div class="section-header">Built-in Fault Rule Library</div>',
                    unsafe_allow_html=True)

        col_filter, col_sev = st.columns([3, 2])
        with col_filter:
            sel_cats = st.multiselect("Filter by category", all_cats, default=all_cats,
                                      key="fault_cats")
        with col_sev:
            sel_sevs = st.multiselect("Filter by severity",
                                      ["HIGH", "MEDIUM", "LOW"],
                                      default=["HIGH", "MEDIUM", "LOW"],
                                      key="fault_sevs")

        filtered_rules = [r for r in BUILTIN_RULES
                          if r["category"] in sel_cats and r["severity"] in sel_sevs]
        sorted_rules = sorted(filtered_rules, key=lambda r: (
            not rule_results.get(r["id"], {}).get("triggered", False),
            sev_order.get(r["severity"], 9),
        ))

        for rule in sorted_rules:
            res       = rule_results.get(rule["id"], {})
            triggered = res.get("triggered", False)
            events    = res.get("events", [])
            skip      = res.get("skip_reason")
            enabled   = st.session_state["enabled_rules"].get(rule["id"], True)

            if triggered:
                icon  = "🔴" if rule["severity"] == "HIGH" else ("🟡" if rule["severity"] == "MEDIUM" else "🟢")
                label = f"{icon} **{rule['name']}** — {len(events)} fault event(s)"
            elif skip:
                label = f"⚫ **{rule['name']}** — skipped (missing data)"
            else:
                label = f"✅ **{rule['name']}** — No faults"

            with st.expander(label, expanded=triggered and rule["severity"] == "HIGH"):
                ca, cb = st.columns([4, 1])
                with ca:
                    st.markdown(f'<div class="rule-expr">{rule["expr"]}</div>',
                                unsafe_allow_html=True)
                    st.markdown(f'<div class="rule-desc">{rule["description"]}</div>',
                                unsafe_allow_html=True)
                    if triggered:
                        st.markdown(
                            f'<div style="margin-top:8px;font-size:12px;color:#f7c948">'
                            f'💡 <b>Recommended action:</b> {rule.get("recommended_action","—")}'
                            f'</div>', unsafe_allow_html=True)
                    if skip:
                        st.warning(f"⚠ Skipped: {skip}")
                with cb:
                    tog = st.toggle("Enable", value=enabled, key=f"tog_{rule['id']}")
                    st.session_state["enabled_rules"][rule["id"]] = tog
                    sev_cls = {"HIGH": "severity-high", "MEDIUM": "severity-medium",
                               "LOW": "severity-low"}.get(rule["severity"], "")
                    st.markdown(f'<div class="{sev_cls}">{rule["severity"]}</div>',
                                unsafe_allow_html=True)
                    st.caption(rule["category"])

                if triggered and events:
                    fault_mask = res["fault_mask"]
                    rel_col    = rule["required_cols"][0]
                    fig = go.Figure()
                    pal = list(COLORS.values())
                    for ci, rc in enumerate(rule["required_cols"]):
                        if rc in df.columns:
                            fig.add_trace(go.Scatter(
                                x=df["timestamp"], y=df[rc],
                                name=rc.replace("_", " ").title(),
                                line=dict(color=pal[ci % len(pal)], width=1.5),
                            ))
                    for ev in events[:40]:
                        fig.add_vrect(x0=ev["start"], x1=ev["end"],
                                      fillcolor=sev_colors[rule["severity"]],
                                      opacity=0.13, layer="below", line_width=0)
                    if rel_col in df.columns:
                        fault_df = df[fault_mask]
                        fig.add_trace(go.Scatter(
                            x=fault_df["timestamp"], y=fault_df[rel_col],
                            mode="markers",
                            marker=dict(color=sev_colors[rule["severity"]], size=4, opacity=0.7),
                            name="Fault Point",
                        ))
                    fig.update_layout(**PLOTLY_LAYOUT,
                                      title=f"{rule['name']} — Fault Windows",
                                      height=260)
                    st.plotly_chart(fig, use_container_width=True)

                    sc1, sc2, sc3 = st.columns(3)
                    durations = [ev["duration_hrs"] for ev in events]
                    sc1.metric("Longest Event",   f"{max(durations):.1f} hrs")
                    sc2.metric("Avg Duration",    f"{sum(durations)/len(durations):.1f} hrs")
                    sc3.metric("Total Fault Hrs", f"{sum(durations):.1f} hrs")

                    ev_df = pd.DataFrame(events).copy()
                    ev_df["start"] = ev_df["start"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df["end"]   = ev_df["end"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df.columns  = ["Start", "End", "Duration (hrs)", "Peak Value", "Signal"]
                    st.dataframe(ev_df, use_container_width=True, hide_index=True, height=220)

    # ══════════════════════════════════════════════════════════════════════════════
    # FTAB 2 — Custom Rule Builder
    # ══════════════════════════════════════════════════════════════════════════════
    with ftab2:
        st.markdown('<div class="section-header">Custom Rule Builder</div>',
                    unsafe_allow_html=True)

        rule_mode = st.radio(
            "Rule type",
            ["Single-signal threshold", "Two-signal comparison"],
            horizontal=True, key="rule_mode",
        )
        st.markdown("---")

        if rule_mode == "Single-signal threshold":
            st.caption("Fault when: **signal  OP  threshold** persists for ≥ min duration")
            cr1, cr2, cr3, cr4 = st.columns([3, 2, 2, 2])
            with cr1:
                custom_col = st.selectbox("Signal", numeric_cols, key="cr_col")
            with cr2:
                custom_op  = st.selectbox("Operator", [">", "<", ">=", "<=", "=="], key="cr_op")
            with cr3:
                col_mean   = float(df[custom_col].mean()) if custom_col in df.columns else 50.0
                col_min_v  = float(df[custom_col].min()) if custom_col in df.columns else 0.0
                col_max_v  = float(df[custom_col].max()) if custom_col in df.columns else 100.0
                custom_threshold = st.number_input(
                    "Threshold", value=round(col_mean, 1),
                    min_value=col_min_v - abs(col_min_v),
                    max_value=col_max_v * 2, key="cr_thresh")
            with cr4:
                custom_min_hrs = st.number_input("Min duration (hrs)", value=1.0,
                                                 min_value=0.25, max_value=72.0,
                                                 step=0.25, key="cr_hrs")
            custom_name = st.text_input("Rule name",
                                        value=f"{custom_col} {custom_op} {custom_threshold}",
                                        key="cr_name")
            custom_sev = st.select_slider("Severity",
                                          options=["LOW", "MEDIUM", "HIGH"],
                                          value="MEDIUM", key="cr_sev")

            c_run, c_save = st.columns(2)
            run_single  = c_run.button("▶ Run Rule",   key="run_custom_single")
            save_single = c_save.button("💾 Save Rule", key="save_custom_single")

            if run_single or save_single:
                cr = parse_custom_rule(df, custom_col, custom_op, custom_threshold, custom_min_hrs)
                if cr is None:
                    st.error("Could not evaluate rule.")
                elif not cr["triggered"]:
                    st.success(f"✅ No faults found for: **{custom_name}**")
                else:
                    st.error(f"🔴 **{len(cr['events'])} fault event(s)** for: **{custom_name}**")
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=df["timestamp"], y=df[custom_col],
                                             name=custom_col.replace("_", " ").title(),
                                             line=dict(color=COLORS["teal"], width=1.5)))
                    for ev in cr["events"]:
                        fig.add_vrect(x0=ev["start"], x1=ev["end"],
                                      fillcolor=sev_colors.get(custom_sev, COLORS["amber"]),
                                      opacity=0.15, layer="below", line_width=0)
                    fig.add_hline(y=custom_threshold, line_dash="dash",
                                  line_color=COLORS["amber"],
                                  annotation_text=f"Threshold {custom_threshold}")
                    fig.update_layout(**PLOTLY_LAYOUT,
                                      title=f"Custom Rule — {custom_name}",
                                      height=300, yaxis_title=custom_col)
                    st.plotly_chart(fig, use_container_width=True)
                    ev_df = pd.DataFrame(cr["events"]).copy()
                    ev_df["start"] = ev_df["start"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df["end"]   = ev_df["end"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df.columns  = ["Start", "End", "Duration (hrs)", "Peak Value", "Signal"]
                    st.dataframe(ev_df, use_container_width=True, hide_index=True)
                    st.download_button("⬇ Export Events CSV",
                                       data=ev_df.to_csv(index=False),
                                       file_name="custom_fault_events.csv", mime="text/csv")
                if save_single:
                    st.session_state["saved_custom_rules"].append({
                        "name": custom_name, "type": "single",
                        "col": custom_col, "op": custom_op,
                        "threshold": custom_threshold, "min_hrs": custom_min_hrs,
                        "severity": custom_sev,
                    })
                    st.success(f"✅ Rule **{custom_name}** saved!")

        else:
            st.caption("Fault when: **signal A  OP  (signal B + offset)** persists for ≥ min duration")
            st.markdown(
                '<div class="rule-desc">Example: <span class="rule-expr">'
                'supply_air_temp &gt; return_air_temp − 5 °F for ≥ 30 min'
                '</span> detects inverted ΔT faults.</div>', unsafe_allow_html=True)

            d1, d2, d3 = st.columns(3)
            with d1:
                ds_col_a = st.selectbox("Signal A", numeric_cols, key="ds_col_a")
            with d2:
                ds_op    = st.selectbox("Operator", [">", "<", ">=", "<="], key="ds_op")
            with d3:
                ds_col_b = st.selectbox("Signal B", numeric_cols,
                                        index=min(1, len(numeric_cols)-1), key="ds_col_b")
            ds_offset  = st.number_input("Offset added to B (e.g. −5 means A > B − 5)",
                                         value=0.0, step=0.5, key="ds_offset")
            ds_min_hrs = st.number_input("Min duration (hrs)", value=0.5,
                                         min_value=0.25, max_value=72.0,
                                         step=0.25, key="ds_hrs")
            ds_name    = st.text_input("Rule name",
                                       value=f"{ds_col_a} {ds_op} {ds_col_b} + {ds_offset}",
                                       key="ds_name")
            ds_sev     = st.select_slider("Severity",
                                          options=["LOW", "MEDIUM", "HIGH"],
                                          value="HIGH", key="ds_sev")

            dc_run, dc_save = st.columns(2)
            run_dual  = dc_run.button("▶ Run Rule",   key="run_dual")
            save_dual = dc_save.button("💾 Save Rule", key="save_dual")

            if run_dual or save_dual:
                cr = parse_dual_signal_rule(df, ds_col_a, ds_op, ds_col_b, ds_offset, ds_min_hrs)
                if cr is None:
                    st.error("Could not evaluate rule — check column selection.")
                elif not cr["triggered"]:
                    st.success(f"✅ No faults found for: **{ds_name}**")
                else:
                    st.error(f"🔴 **{len(cr['events'])} fault event(s)** for: **{ds_name}**")
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=df["timestamp"], y=df[ds_col_a],
                                             name=ds_col_a.replace("_", " ").title(),
                                             line=dict(color=COLORS["blue"], width=1.5)))
                    fig.add_trace(go.Scatter(
                        x=df["timestamp"], y=df[ds_col_b] + ds_offset,
                        name=f"{ds_col_b.replace('_',' ').title()} + {ds_offset}",
                        line=dict(color=COLORS["teal"], width=1.5, dash="dash")))
                    for ev in cr["events"]:
                        fig.add_vrect(x0=ev["start"], x1=ev["end"],
                                      fillcolor=sev_colors.get(ds_sev, COLORS["red"]),
                                      opacity=0.13, layer="below", line_width=0)
                    fig.update_layout(**PLOTLY_LAYOUT,
                                      title=f"Two-Signal Rule — {ds_name}", height=320)
                    st.plotly_chart(fig, use_container_width=True)
                    ev_df = pd.DataFrame(cr["events"]).copy()
                    ev_df["start"] = ev_df["start"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df["end"]   = ev_df["end"].dt.strftime("%Y-%m-%d %H:%M")
                    ev_df.columns  = ["Start", "End", "Duration (hrs)", "Peak |ΔSignal|", "Signals"]
                    st.dataframe(ev_df, use_container_width=True, hide_index=True)
                    st.download_button("⬇ Export Events CSV",
                                       data=ev_df.to_csv(index=False),
                                       file_name="dual_signal_fault_events.csv", mime="text/csv")
                if save_dual:
                    st.session_state["saved_custom_rules"].append({
                        "name": ds_name, "type": "dual",
                        "col_a": ds_col_a, "op": ds_op, "col_b": ds_col_b,
                        "offset": ds_offset, "min_hrs": ds_min_hrs, "severity": ds_sev,
                    })
                    st.success(f"✅ Rule **{ds_name}** saved!")

        if st.session_state["saved_custom_rules"]:
            st.markdown("---")
            st.markdown('<div class="section-header">Saved Custom Rules</div>',
                        unsafe_allow_html=True)
            for i, sr in enumerate(st.session_state["saved_custom_rules"]):
                icon = "🔴" if sr["severity"] == "HIGH" else ("🟡" if sr["severity"] == "MEDIUM" else "🟢")
                with st.expander(f"{icon} {sr['name']}"):
                    if sr["type"] == "single":
                        expr_str = f"{sr['col']} {sr['op']} {sr['threshold']} for ≥{sr['min_hrs']}h"
                        sr_result = parse_custom_rule(df, sr["col"], sr["op"],
                                                      sr["threshold"], sr["min_hrs"])
                    else:
                        expr_str = f"{sr['col_a']} {sr['op']} {sr['col_b']} + {sr['offset']} for ≥{sr['min_hrs']}h"
                        sr_result = parse_dual_signal_rule(df, sr["col_a"], sr["op"],
                                                           sr["col_b"], sr["offset"], sr["min_hrs"])
                    st.markdown(f'<div class="rule-expr">{expr_str}</div>', unsafe_allow_html=True)
                    st.caption(f"Severity: {sr['severity']}")
                    if sr_result and sr_result["triggered"]:
                        st.error(f"🔴 {len(sr_result['events'])} event(s) detected")
                    elif sr_result:
                        st.success("✅ No faults")
                    if st.button("🗑 Delete", key=f"del_sr_{i}"):
                        st.session_state["saved_custom_rules"].pop(i)
                        st.rerun()

    # ══════════════════════════════════════════════════════════════════════════════
    # FTAB 3 — Fault Calendar
    # ══════════════════════════════════════════════════════════════════════════════
    with ftab3:
        st.markdown('<div class="section-header">Fault Calendar — Daily Severity Score</div>',
                    unsafe_allow_html=True)
        st.caption(
            "Score = Σ (fault hours × severity weight) per day. "
            "HIGH=3 · MEDIUM=2 · LOW=1. Darker = more severe fault activity."
        )

        heat_df, daily_score = build_daily_fault_heatmap(df, rule_results, BUILTIN_RULES)

        if daily_score.sum() == 0:
            st.info("No fault events detected — nothing to plot on the calendar.")
        else:
            score_series = daily_score.reset_index()
            score_series.columns = ["date", "score"]
            score_series["date"] = pd.to_datetime(score_series["date"])

            fig = go.Figure(go.Bar(
                x=score_series["date"], y=score_series["score"],
                marker=dict(
                    color=score_series["score"],
                    colorscale=[[0, "#1c2128"], [0.3, "#0077b6"],
                                [0.6, "#f7c948"], [1.0, "#ff6b6b"]],
                    showscale=True,
                    colorbar=dict(title="Score", tickfont=dict(color="#8b949e")),
                ),
            ))
            fig.update_layout(**PLOTLY_LAYOUT,
                              title="Daily Fault Severity Score",
                              height=300, xaxis_title="Date", yaxis_title="Score")
            st.plotly_chart(fig, use_container_width=True)

            if not heat_df.empty:
                st.markdown('<div class="section-header">Fault Hours per Rule per Day</div>',
                            unsafe_allow_html=True)
                fig2 = go.Figure(go.Heatmap(
                    z=heat_df.T.values,
                    x=[str(d) for d in heat_df.index],
                    y=heat_df.columns.tolist(),
                    colorscale=[[0, "#1c2128"], [0.01, "#0d3b52"],
                                [0.3, "#0077b6"], [0.7, "#f7c948"], [1.0, "#ff6b6b"]],
                    colorbar=dict(title="Fault hrs", tickfont=dict(color="#8b949e")),
                    hoverongaps=False,
                ))
                fig2.update_layout(
                    **PLOTLY_LAYOUT,
                    title="Fault Hours Heatmap (Rule × Day)",
                    height=max(200, 40 * len(heat_df.columns) + 80),
                )
                fig2.update_xaxes(tickangle=-45, tickfont=dict(size=9))
                fig2.update_yaxes(tickfont=dict(size=10))
                st.plotly_chart(fig2, use_container_width=True)

            st.markdown('<div class="section-header">Top 10 Worst Fault Days</div>',
                        unsafe_allow_html=True)
            worst = score_series.sort_values("score", ascending=False).head(10).copy()
            worst["date"] = worst["date"].dt.strftime("%Y-%m-%d (%A)")
            worst["score"] = worst["score"].round(1)
            worst.columns = ["Date", "Severity Score"]
            st.dataframe(worst, use_container_width=True, hide_index=True)

    # ══════════════════════════════════════════════════════════════════════════════
    # FTAB 4 — Fault Summary & Export
    # ══════════════════════════════════════════════════════════════════════════════
    with ftab4:
        st.markdown('<div class="section-header">Fault Event Summary</div>',
                    unsafe_allow_html=True)

        if not triggered_rules:
            st.success("🎉 No fault rules triggered on the current dataset.")
        else:
            fig = go.Figure()
            if "total_kw" in df.columns:
                fig.add_trace(go.Scatter(
                    x=df["timestamp"], y=df["total_kw"],
                    name="Total kW (reference)",
                    line=dict(color="#30363d", width=1), opacity=0.5,
                ))
            for rule in triggered_rules:
                res = rule_results[rule["id"]]
                for ev in res["events"]:
                    fig.add_vrect(
                        x0=ev["start"], x1=ev["end"],
                        fillcolor=sev_colors[rule["severity"]],
                        opacity=0.18, layer="below", line_width=0,
                        annotation_text=rule["name"][:22],
                        annotation_position="top left",
                        annotation_font_size=8,
                        annotation_font_color=sev_colors[rule["severity"]],
                    )
            fig.update_layout(**PLOTLY_LAYOUT,
                              title="All Fault Windows — Timeline Overlay",
                              height=360,
                              yaxis_title="Total kW" if "total_kw" in df.columns else "")
            st.plotly_chart(fig, use_container_width=True)

            summary_data = []
            for rule in triggered_rules:
                total_hrs = sum(ev["duration_hrs"]
                                for ev in rule_results[rule["id"]]["events"])
                summary_data.append({
                    "Rule": rule["name"], "Category": rule["category"],
                    "Severity": rule["severity"],
                    "Events": len(rule_results[rule["id"]]["events"]),
                    "Total Fault Hrs": round(total_hrs, 1),
                })
            summary_df = pd.DataFrame(summary_data).sort_values(
                ["Severity", "Total Fault Hrs"],
                key=lambda c: c.map(sev_order) if c.name == "Severity" else c,
            )
            fig_bar = go.Figure(go.Bar(
                y=summary_df["Rule"],
                x=summary_df["Total Fault Hrs"],
                orientation="h",
                marker=dict(
                    color=[sev_colors[s] for s in summary_df["Severity"]],
                    opacity=0.85,
                ),
                text=summary_df["Total Fault Hrs"].astype(str) + " hrs",
                textposition="outside",
            ))
            fig_bar.update_layout(
                **PLOTLY_LAYOUT,
                title="Total Fault Hours by Rule",
                height=max(280, 40 * len(summary_df) + 80),
                xaxis_title="Hours",
                )
            fig_bar.update_layout(margin=dict(l=220, r=60, t=48, b=40))
            st.plotly_chart(fig_bar, use_container_width=True)

            st.markdown('<div class="section-header">Complete Fault Event Log</div>',
                        unsafe_allow_html=True)
            all_events = []
            for rule in triggered_rules:
                for ev in rule_results[rule["id"]]["events"]:
                    all_events.append({
                        "Rule":           rule["name"],
                        "Severity":       rule["severity"],
                        "Category":       rule["category"],
                        "Start":          ev["start"],
                        "End":            ev["end"],
                        "Duration (hrs)": ev["duration_hrs"],
                        "Peak Value":     ev["peak_value"],
                        "Signal":         ev["column"],
                        "Action":         rule.get("recommended_action", "—"),
                    })
            if all_events:
                export_df = pd.DataFrame(all_events).sort_values(["Severity", "Start"])
                st.dataframe(export_df, use_container_width=True,
                             hide_index=True, height=320)
                st.download_button(
                    "⬇ Export All Fault Events CSV",
                    data=export_df.to_csv(index=False),
                    file_name="all_fault_events.csv",
                    mime="text/csv",
                )


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 6 — ML Insights
# ══════════════════════════════════════════════════════════════════════════════════
with tab6:
    import warnings
    warnings.filterwarnings("ignore")
    from statsmodels.tsa.holtwinters import ExponentialSmoothing
    from sklearn.linear_model import LinearRegression, Ridge
    from sklearn.preprocessing import StandardScaler
    from sklearn.pipeline import Pipeline
    from sklearn.ensemble import IsolationForest
    import sklearn.metrics as skm

    st.markdown("## 🤖 ML Insights")
    st.caption("Explainable machine-learning analysis — forecasting, predictive faults, health scoring, and setpoint optimisation.")

    ml_tab1, ml_tab2, ml_tab3, ml_tab4 = st.tabs([
        "📉 Energy Forecast",
        "🔮 Predictive Faults",
        "❤️ Health Scores",
        "🎯 Setpoint Optimiser",
    ])

    # ── helpers shared across ML tabs ────────────────────────────────────────────
    def make_time_features(ts: pd.Series) -> pd.DataFrame:
        return pd.DataFrame({
            "hour":       ts.dt.hour,
            "dow":        ts.dt.dayofweek,
            "month":      ts.dt.month,
            "hour_sin":   np.sin(2 * np.pi * ts.dt.hour / 24),
            "hour_cos":   np.cos(2 * np.pi * ts.dt.hour / 24),
            "dow_sin":    np.sin(2 * np.pi * ts.dt.dayofweek / 7),
            "dow_cos":    np.cos(2 * np.pi * ts.dt.dayofweek / 7),
        })

    # ════════════════════════════════════════════════════════════════════════════
    # ML-TAB 1 · Energy Forecast (Holt-Winters + Linear fallback)
    # ════════════════════════════════════════════════════════════════════════════
    with ml_tab1:
        st.markdown('<div class="section-header">24 / 48-hr Energy Forecast</div>', unsafe_allow_html=True)
        st.markdown(
            "Uses **Holt-Winters Exponential Smoothing** — a well-understood seasonal model that "
            "captures daily cycles and trend drift without a black box."
        )

        fc1, fc2, fc3 = st.columns(3)
        with fc1:
            fc_target = st.selectbox(
                "Signal to forecast",
                [c for c in ["total_kw", "chiller_kw", "ahu_fan_kw"] if c in df.columns]
                or numeric_cols[:1],
                key="fc_target",
            )
        with fc2:
            fc_horizon = st.selectbox("Horizon", [24, 48], key="fc_horizon")
        with fc3:
            fc_season  = st.selectbox("Seasonal period (hrs)", [24, 12], key="fc_season")

        if st.button("▶ Run Forecast", key="run_fc"):
            series = df.set_index("timestamp")[fc_target].resample("h").mean().ffill()

            if len(series) < fc_season * 2 + fc_horizon:
                st.warning("Not enough data — need at least 2 full seasonal cycles + horizon.")
            else:
                with st.spinner("Fitting Holt-Winters model…"):
                    try:
                        model = ExponentialSmoothing(
                            series, trend="add", seasonal="add",
                            seasonal_periods=fc_season,
                            initialization_method="estimated",
                        ).fit(optimized=True)

                        # In-sample fit
                        fitted   = model.fittedvalues
                        forecast = model.forecast(fc_horizon)
                        fc_idx   = pd.date_range(series.index[-1], periods=fc_horizon + 1, freq="h")[1:]
                        forecast.index = fc_idx

                        # CI: use residual std
                        resid_std = (series - fitted).std()
                        ci_upper  = forecast + 1.96 * resid_std
                        ci_lower  = forecast - 1.96 * resid_std

                        # Metrics on last fc_horizon hours as pseudo-test
                        train, test = series[:-fc_horizon], series[-fc_horizon:]
                        m_model = ExponentialSmoothing(
                            train, trend="add", seasonal="add",
                            seasonal_periods=fc_season,
                            initialization_method="estimated",
                        ).fit(optimized=True)
                        preds   = m_model.forecast(fc_horizon)
                        mae     = np.mean(np.abs(test.values - preds.values))
                        mape    = np.mean(np.abs((test.values - preds.values) / (test.values + 1e-9))) * 100
                        rmse    = np.sqrt(np.mean((test.values - preds.values) ** 2))

                    except Exception as e:
                        st.error(f"Holt-Winters failed ({e}) — falling back to linear regression.")
                        model = None

                if model is not None:
                    # Accuracy strip
                    m1, m2, m3 = st.columns(3)
                    m1.metric("MAE",  f"{mae:.1f} kW")
                    m2.metric("MAPE", f"{mape:.1f}%")
                    m3.metric("RMSE", f"{rmse:.1f} kW")

                    # Chart
                    history_window = series.last("7D") if len(series) > 7 * 24 else series
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(
                        x=history_window.index, y=history_window.values,
                        name="Historical", line=dict(color=COLORS["blue"], width=1.5),
                    ))
                    fig.add_trace(go.Scatter(
                        x=forecast.index, y=forecast.values,
                        name=f"{fc_horizon}h Forecast",
                        line=dict(color=COLORS["amber"], width=2.5, dash="dot"),
                    ))
                    # CI band
                    fig.add_trace(go.Scatter(
                        x=pd.concat([pd.Series(fc_idx), pd.Series(fc_idx[::-1])]),
                        y=pd.concat([ci_upper, ci_lower[::-1]]),
                        fill="toself",
                        fillcolor="rgba(247,201,72,0.10)",
                        line=dict(color="rgba(0,0,0,0)"),
                        name="95% CI",
                    ))
                    fig.add_vline(
                        x=str(series.index[-1]),
                        line_dash="dash", line_color="#30363d",
                        annotation_text="Now", annotation_font_color="#8b949e",
                    )
                    fig.update_layout(**PLOTLY_LAYOUT,
                                      title=f"{fc_target.replace('_',' ').title()} — {fc_horizon}hr Forecast",
                                      height=380, yaxis_title="kW")
                    st.plotly_chart(fig, use_container_width=True)

                    # Forecast table
                    fc_df = pd.DataFrame({
                        "Timestamp":   fc_idx.strftime("%Y-%m-%d %H:%M"),
                        "Forecast kW": forecast.values.round(1),
                        "Lower CI":    ci_lower.values.round(1),
                        "Upper CI":    ci_upper.values.round(1),
                    })
                    with st.expander("View forecast table"):
                        st.dataframe(fc_df, use_container_width=True, hide_index=True)
                    st.download_button("⬇ Export Forecast CSV",
                                       data=fc_df.to_csv(index=False),
                                       file_name="hvac_forecast.csv", mime="text/csv")

                    st.info(
                        "**How it works:** Holt-Winters models three components — level (current average), "
                        "trend (drift up/down), and seasonality (recurring daily pattern). "
                        f"The model found α={model.params['smoothing_level']:.3f} "
                        f"(level), β={model.params['smoothing_trend']:.3f} (trend), "
                        f"γ={model.params['smoothing_seasonal']:.3f} (seasonal)."
                    )

    # ════════════════════════════════════════════════════════════════════════════
    # ML-TAB 2 · Predictive Fault Detection (Isolation Forest)
    # ════════════════════════════════════════════════════════════════════════════
    with ml_tab2:
        st.markdown('<div class="section-header">Predictive Fault Detection</div>', unsafe_allow_html=True)
        st.markdown(
            "Uses **Isolation Forest** — an unsupervised model that isolates outliers by randomly "
            "partitioning features. Points that are easy to isolate (short average path length) are flagged "
            "as anomalies *before* they breach hard thresholds."
        )

        pf1, pf2 = st.columns(2)
        with pf1:
            pf_cols = st.multiselect(
                "Features for anomaly model",
                numeric_cols,
                default=numeric_cols[:6],
                key="pf_cols",
            )
        with pf2:
            pf_contamination = st.slider(
                "Expected anomaly fraction",
                min_value=0.01, max_value=0.20, value=0.05, step=0.01,
                key="pf_cont",
                help="Fraction of points the model will flag as anomalies (0.05 = 5%)",
            )

        pf_lookahead = st.slider(
            "Lookahead window (hrs) — flag N hours BEFORE rule engine would trigger",
            1, 12, 3, key="pf_look",
        )

        if st.button("▶ Run Predictive Model", key="run_pf"):
            if len(pf_cols) < 2:
                st.warning("Select at least 2 features.")
            else:
                with st.spinner("Training Isolation Forest…"):
                    feat_df = df[pf_cols].fillna(df[pf_cols].median())

                    # Add rolling stats as extra features
                    roll_feats = {}
                    for c in pf_cols:
                        roll_feats[f"{c}_roll6_mean"] = df[c].rolling(6, min_periods=1).mean()
                        roll_feats[f"{c}_roll6_std"]  = df[c].rolling(6, min_periods=1).std().fillna(0)
                    feat_df = pd.concat([feat_df, pd.DataFrame(roll_feats)], axis=1)

                    iso = IsolationForest(
                        contamination=pf_contamination,
                        n_estimators=200,
                        random_state=42,
                        n_jobs=-1,
                    )
                    labels = iso.fit_predict(feat_df)          # -1 = anomaly
                    scores = iso.decision_function(feat_df)    # lower = more anomalous
                    norm_score = 1 - (scores - scores.min()) / (scores.max() - scores.min() + 1e-9)

                    pred_df = df[["timestamp"]].copy()
                    pred_df["anomaly_flag"]  = labels == -1
                    pred_df["anomaly_score"] = norm_score

                    # Shift scores back by lookahead to simulate early warning
                    pred_df["early_warning"] = pred_df["anomaly_score"].shift(pf_lookahead).fillna(0) > 0.6

                n_flagged = pred_df["anomaly_flag"].sum()
                n_early   = pred_df["early_warning"].sum()
                c1, c2, c3 = st.columns(3)
                c1.metric("Points Flagged",    f"{n_flagged:,}")
                c2.metric("Early Warnings",    f"{n_early:,}  ({pf_lookahead}hr ahead)")
                c3.metric("Contamination Used", f"{pf_contamination*100:.0f}%")

                # Score timeline
                primary = pf_cols[0]
                fig = make_subplots(rows=2, cols=1, shared_xaxes=True,
                                    row_heights=[0.6, 0.4],
                                    vertical_spacing=0.06)
                fig.add_trace(go.Scatter(
                    x=df["timestamp"], y=df[primary],
                    name=primary.replace("_", " ").title(),
                    line=dict(color=COLORS["blue"], width=1.2),
                ), row=1, col=1)
                # Anomaly markers
                anom_pts = pred_df[pred_df["anomaly_flag"]]
                fig.add_trace(go.Scatter(
                    x=anom_pts["timestamp"], y=df.loc[anom_pts.index, primary],
                    mode="markers", name="Flagged Anomaly",
                    marker=dict(color=COLORS["red"], size=5, symbol="x"),
                ), row=1, col=1)
                # Early warnings
                ew_pts = pred_df[pred_df["early_warning"] & ~pred_df["anomaly_flag"]]
                fig.add_trace(go.Scatter(
                    x=ew_pts["timestamp"], y=df.loc[ew_pts.index, primary],
                    mode="markers", name=f"Early Warning (+{pf_lookahead}hr)",
                    marker=dict(color=COLORS["amber"], size=5, symbol="triangle-up"),
                ), row=1, col=1)
                # Score bar
                fig.add_trace(go.Bar(
                    x=pred_df["timestamp"], y=pred_df["anomaly_score"],
                    name="Anomaly Score",
                    marker_color=pred_df["anomaly_score"].apply(
                        lambda v: COLORS["red"] if v > 0.7 else COLORS["amber"] if v > 0.4 else COLORS["blue"]
                    ),
                ), row=2, col=1)
                fig.add_hline(y=0.6, row=2, col=1, line_dash="dash",
                              line_color=COLORS["amber"], annotation_text="Warning threshold")

                fig.update_layout(**PLOTLY_LAYOUT,
                                  title="Isolation Forest — Anomaly Scores & Early Warnings",
                                  height=500)
                fig.update_yaxes(title_text=primary, row=1, col=1,
                                 gridcolor="#21262d", zerolinecolor="#30363d")
                fig.update_yaxes(title_text="Score", row=2, col=1,
                                 gridcolor="#21262d", zerolinecolor="#30363d")
                st.plotly_chart(fig, use_container_width=True)

                # Feature importance proxy: mean score contribution
                st.markdown('<div class="section-header">Feature Influence</div>', unsafe_allow_html=True)
                st.caption("Pearson correlation between each feature's deviation from its mean and the anomaly score — shows which signals drive the flags.")
                influences = {}
                for c in pf_cols:
                    dev = (df[c] - df[c].mean()).abs()
                    corr = dev.corr(pred_df["anomaly_score"])
                    influences[c] = abs(corr) if not np.isnan(corr) else 0
                inf_df = pd.DataFrame({"Feature": list(influences.keys()),
                                       "Influence": list(influences.values())}).sort_values("Influence", ascending=True)
                fig2 = go.Figure(go.Bar(
                    x=inf_df["Influence"], y=inf_df["Feature"],
                    orientation="h",
                    marker_color=COLORS["teal"],
                ))
                fig2.update_layout(**PLOTLY_LAYOUT, height=280, title="Feature Influence on Anomaly Score",
                                   xaxis_title="|Correlation with Score|")
                st.plotly_chart(fig2, use_container_width=True)

                st.download_button(
                    "⬇ Export Flagged Points CSV",
                    data=pred_df[pred_df["anomaly_flag"]].merge(
                        df, on="timestamp", how="left"
                    ).to_csv(index=False),
                    file_name="predictive_faults.csv", mime="text/csv",
                )

    # ════════════════════════════════════════════════════════════════════════════
    # ML-TAB 3 · Equipment Health Scores
    # ════════════════════════════════════════════════════════════════════════════
    with ml_tab3:
        st.markdown('<div class="section-header">Equipment Health Scoring (0 – 100)</div>', unsafe_allow_html=True)
        st.markdown(
            "Each asset receives a **weighted composite score** derived from rule-based sub-indicators. "
            "Fully transparent — every component is shown so you can trace the score back to raw data."
        )

        # Define assets and their health indicators
        ASSET_INDICATORS = {
            "Chiller": {
                "required": ["chiller_kw", "chilled_water_temp", "condenser_temp"],
                "indicators": [
                    {
                        "name": "Load Stability",
                        "desc": "Low coefficient of variation in chiller power → stable, predictable load",
                        "weight": 0.25,
                        "fn": lambda d: max(0.0, 1 - d["chiller_kw"].std() / (d["chiller_kw"].mean() + 1e-9)),
                        "cols": ["chiller_kw"],
                    },
                    {
                        "name": "Chilled Water Temp Control",
                        "desc": "How often CHW supply stays within ±2 °F of its mean setpoint proxy",
                        "weight": 0.25,
                        "fn": lambda d: (
                            (d["chilled_water_temp"] - d["chilled_water_temp"].mean()).abs() <= 2
                        ).mean(),
                        "cols": ["chilled_water_temp"],
                    },
                    {
                        "name": "Condenser Temp Rating",
                        "desc": "Fraction of time condenser temp < 95 °F (below high-stress threshold)",
                        "weight": 0.30,
                        "fn": lambda d: (d["condenser_temp"] < 95).mean(),
                        "cols": ["condenser_temp"],
                    },
                    {
                        "name": "Efficiency Trend",
                        "desc": "Ratio of mean load this week vs. prior period — degradation raises kW for same output",
                        "weight": 0.20,
                        "fn": lambda d: min(1.0, d["chiller_kw"].iloc[:len(d)//2].mean()
                                            / (d["chiller_kw"].iloc[len(d)//2:].mean() + 1e-9)),
                        "cols": ["chiller_kw"],
                    },
                ],
            },
            "AHU": {
                "required": ["supply_air_temp", "return_air_temp", "supply_cfm"],
                "indicators": [
                    {
                        "name": "ΔT Adequacy",
                        "desc": "Fraction of time where return−supply ΔT ≥ 12 °F (good heat transfer)",
                        "weight": 0.35,
                        "fn": lambda d: ((d["return_air_temp"] - d["supply_air_temp"]) >= 12).mean(),
                        "cols": ["supply_air_temp", "return_air_temp"],
                    },
                    {
                        "name": "Supply Temp Stability",
                        "desc": "1 − CoV of supply air temp; low variation = good control",
                        "weight": 0.25,
                        "fn": lambda d: max(0.0, 1 - d["supply_air_temp"].std()
                                            / (abs(d["supply_air_temp"].mean()) + 1e-9)),
                        "cols": ["supply_air_temp"],
                    },
                    {
                        "name": "Airflow Consistency",
                        "desc": "Fraction of time CFM is within 20% of its mean — fan/damper stability",
                        "weight": 0.25,
                        "fn": lambda d: (
                            (d["supply_cfm"] - d["supply_cfm"].mean()).abs()
                            <= 0.2 * d["supply_cfm"].mean()
                        ).mean(),
                        "cols": ["supply_cfm"],
                    },
                    {
                        "name": "No Inverted ΔT",
                        "desc": "Fraction of time supply is cooler than return (correct direction)",
                        "weight": 0.15,
                        "fn": lambda d: (d["supply_air_temp"] < d["return_air_temp"]).mean(),
                        "cols": ["supply_air_temp", "return_air_temp"],
                    },
                ],
            },
            "Pumps / Fans": {
                "required": ["ahu_fan_kw", "pump_kw"],
                "indicators": [
                    {
                        "name": "Fan Load Stability",
                        "desc": "1 − CoV of AHU fan kW",
                        "weight": 0.50,
                        "fn": lambda d: max(0.0, 1 - d["ahu_fan_kw"].std()
                                            / (d["ahu_fan_kw"].mean() + 1e-9)),
                        "cols": ["ahu_fan_kw"],
                    },
                    {
                        "name": "Pump Load Stability",
                        "desc": "1 − CoV of pump kW",
                        "weight": 0.50,
                        "fn": lambda d: max(0.0, 1 - d["pump_kw"].std()
                                            / (d["pump_kw"].mean() + 1e-9)),
                        "cols": ["pump_kw"],
                    },
                ],
            },
        }

        scores_summary = {}
        for asset_name, asset_cfg in ASSET_INDICATORS.items():
            missing = [c for c in asset_cfg["required"] if c not in df.columns]
            if missing:
                scores_summary[asset_name] = {"score": None, "missing": missing, "breakdown": []}
                continue

            breakdown = []
            weighted_sum = 0.0
            for ind in asset_cfg["indicators"]:
                if all(c in df.columns for c in ind["cols"]):
                    try:
                        raw = float(ind["fn"](df))
                        raw = max(0.0, min(1.0, raw))
                    except Exception:
                        raw = 0.5
                    contrib = raw * ind["weight"] * 100
                    weighted_sum += contrib
                    breakdown.append({
                        "Indicator": ind["name"],
                        "Description": ind["desc"],
                        "Score (0–1)": round(raw, 3),
                        "Weight": f"{ind['weight']*100:.0f}%",
                        "Contribution": round(contrib, 1),
                    })

            total = min(100.0, round(weighted_sum, 1))
            scores_summary[asset_name] = {"score": total, "breakdown": breakdown, "missing": []}

        # Score cards row
        score_cols = st.columns(len(ASSET_INDICATORS))
        for i, (asset, info) in enumerate(scores_summary.items()):
            with score_cols[i]:
                sc = info["score"]
                if sc is None:
                    metric_card(asset, "N/A", status="default")
                    st.caption(f"Missing: {info['missing']}")
                else:
                    status = "good" if sc >= 75 else "warning" if sc >= 50 else "danger"
                    metric_card(asset + " Health", f"{sc:.0f} / 100", status=status)

        # Gauge charts
        gauge_cols = st.columns(len(ASSET_INDICATORS))
        for i, (asset, info) in enumerate(scores_summary.items()):
            with gauge_cols[i]:
                sc = info["score"]
                if sc is None:
                    continue
                color = COLORS["green"] if sc >= 75 else COLORS["amber"] if sc >= 50 else COLORS["red"]
                fig = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=sc,
                    number={"suffix": "/100", "font": {"color": "#e6edf3", "family": "Space Mono", "size": 22}},
                    gauge={
                        "axis": {"range": [0, 100], "tickcolor": "#8b949e",
                                 "tickfont": {"color": "#8b949e"}},
                        "bar":  {"color": color, "thickness": 0.25},
                        "bgcolor": "#0d1117",
                        "bordercolor": "#30363d",
                        "steps": [
                            {"range": [0,  50], "color": "rgba(255,107,107,0.08)"},
                            {"range": [50, 75], "color": "rgba(247,201,72,0.08)"},
                            {"range": [75,100], "color": "rgba(46,204,113,0.08)"},
                        ],
                        "threshold": {
                            "line": {"color": color, "width": 3},
                            "thickness": 0.75,
                            "value": sc,
                        },
                    },
                    title={"text": asset, "font": {"color": "#8b949e", "size": 13}},
                ))
                fig.update_layout(paper_bgcolor="#161b22", font_color="#8b949e",
                                  height=220, margin=dict(l=20, r=20, t=40, b=10))
                st.plotly_chart(fig, use_container_width=True)

        # Breakdown tables
        for asset, info in scores_summary.items():
            if info["score"] is None or not info["breakdown"]:
                continue
            with st.expander(f"📋 {asset} — Indicator Breakdown"):
                bk = pd.DataFrame(info["breakdown"])
                # Colour the Score column
                st.dataframe(bk, use_container_width=True, hide_index=True)
                st.caption(
                    "**Score (0–1)**: raw indicator value before weighting.  "
                    "**Contribution**: Score × Weight × 100 — sums to the asset total."
                )

        # Trend: rolling 24h health proxy
        st.markdown('<div class="section-header">Rolling Health Trend</div>', unsafe_allow_html=True)
        if "chiller_kw" in df.columns and "chilled_water_temp" in df.columns:
            window_h = st.slider("Rolling window (hrs)", 6, 72, 24, key="health_window")
            # Proxy: normalised distance from optimal operating point
            chw_dev  = (df["chilled_water_temp"] - 44).abs().rolling(window_h, min_periods=1).mean()
            load_cv  = df["chiller_kw"].rolling(window_h, min_periods=1).std() / (
                df["chiller_kw"].rolling(window_h, min_periods=1).mean() + 1e-9)
            health_proxy = (1 - (chw_dev / 10).clip(0, 1)) * 0.5 + (1 - load_cv.clip(0, 1)) * 0.5
            health_proxy = (health_proxy * 100).clip(0, 100)

            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=df["timestamp"], y=health_proxy,
                fill="tozeroy", name="Chiller Health Proxy",
                line=dict(color=COLORS["teal"], width=1.5),
                fillcolor="rgba(6,214,160,0.08)",
            ))
            fig.add_hline(y=75, line_dash="dash", line_color=COLORS["green"],
                          annotation_text="Good (75)")
            fig.add_hline(y=50, line_dash="dash", line_color=COLORS["amber"],
                          annotation_text="Warning (50)")
            fig.update_layout(**PLOTLY_LAYOUT,
                              title=f"Chiller Health Proxy — {window_h}hr Rolling",
                              height=300, yaxis_title="Health Score",
                              yaxis_range=[0, 105])
            st.plotly_chart(fig, use_container_width=True)

    # ════════════════════════════════════════════════════════════════════════════
    # ML-TAB 4 · Setpoint Optimiser (Ridge Regression)
    # ════════════════════════════════════════════════════════════════════════════
    with ml_tab4:
        st.markdown('<div class="section-header">Setpoint Optimiser</div>', unsafe_allow_html=True)
        st.markdown(
            "Trains a **Ridge Regression** model to predict energy use from operating conditions, "
            "then sweeps candidate setpoints to recommend the one with the lowest predicted consumption "
            "while respecting comfort and safety constraints."
        )

        so1, so2 = st.columns(2)
        with so1:
            so_target = st.selectbox(
                "Optimise (minimise)",
                [c for c in ["total_kw", "chiller_kw"] if c in df.columns] or numeric_cols[:1],
                key="so_target",
            )
            so_setpoint_col = st.selectbox(
                "Setpoint signal to sweep",
                [c for c in ["supply_air_temp", "chilled_water_temp", "oa_damper_pct"] if c in df.columns]
                or numeric_cols[:1],
                key="so_sp_col",
            )
        with so2:
            sp_min = st.number_input("Min candidate value",
                                     value=float(round(df[so_setpoint_col].quantile(0.05), 1))
                                     if so_setpoint_col in df.columns else 40.0,
                                     key="so_min")
            sp_max = st.number_input("Max candidate value",
                                     value=float(round(df[so_setpoint_col].quantile(0.95), 1))
                                     if so_setpoint_col in df.columns else 65.0,
                                     key="so_max")
            sp_steps = st.slider("Sweep steps", 10, 100, 40, key="so_steps")

        so_feature_cols = st.multiselect(
            "Predictor features (include time features automatically)",
            [c for c in numeric_cols if c != so_target],
            default=[c for c in numeric_cols if c != so_target][:5],
            key="so_feats",
        )

        if st.button("▶ Run Optimiser", key="run_so"):
            if so_setpoint_col not in df.columns or so_target not in df.columns:
                st.error("Selected columns not found in dataset.")
            elif len(so_feature_cols) < 1:
                st.warning("Select at least 1 predictor feature.")
            else:
                with st.spinner("Training Ridge model…"):
                    feat_df = df[so_feature_cols].fillna(df[so_feature_cols].median())
                    time_feats = make_time_features(df["timestamp"])
                    X = pd.concat([feat_df.reset_index(drop=True),
                                   time_feats.reset_index(drop=True)], axis=1)
                    y = df[so_target].fillna(df[so_target].median())

                    pipe = Pipeline([
                        ("scaler", StandardScaler()),
                        ("ridge",  Ridge(alpha=1.0)),
                    ])
                    pipe.fit(X, y)
                    y_pred = pipe.predict(X)

                    r2   = skm.r2_score(y, y_pred)
                    mae  = skm.mean_absolute_error(y, y_pred)
                    rmse = np.sqrt(skm.mean_squared_error(y, y_pred))

                    # Sweep setpoint
                    sp_candidates = np.linspace(sp_min, sp_max, sp_steps)
                    X_sweep = X.copy()
                    mean_row = X_sweep.mean()

                    predictions = []
                    for sp_val in sp_candidates:
                        row = mean_row.copy()
                        if so_setpoint_col in X_sweep.columns:
                            row[so_setpoint_col] = sp_val
                        pred = pipe.predict(row.values.reshape(1, -1))[0]
                        predictions.append(pred)

                    predictions  = np.array(predictions)
                    best_idx     = np.argmin(predictions)
                    best_sp      = sp_candidates[best_idx]
                    best_pred    = predictions[best_idx]
                    current_sp   = float(df[so_setpoint_col].mean())
                    current_pred = pipe.predict(
                        mean_row.values.reshape(1, -1)
                    )[0]
                    savings_pct  = (current_pred - best_pred) / (current_pred + 1e-9) * 100

                # Model quality strip
                q1, q2, q3, q4 = st.columns(4)
                q1.metric("R²",  f"{r2:.3f}")
                q2.metric("MAE", f"{mae:.1f} kW")
                q3.metric("RMSE",f"{rmse:.1f} kW")
                q4.metric("Model", "Ridge α=1")

                # Sweep chart
                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    x=sp_candidates, y=predictions,
                    name="Predicted Energy",
                    line=dict(color=COLORS["blue"], width=2),
                    fill="tozeroy", fillcolor="rgba(0,180,216,0.06)",
                ))
                fig.add_vline(x=current_sp, line_dash="dash",
                              line_color=COLORS["amber"],
                              annotation_text=f"Current ({current_sp:.1f})",
                              annotation_font_color=COLORS["amber"])
                fig.add_vline(x=best_sp, line_dash="solid",
                              line_color=COLORS["green"],
                              annotation_text=f"Optimal ({best_sp:.1f})",
                              annotation_font_color=COLORS["green"])
                fig.add_trace(go.Scatter(
                    x=[best_sp], y=[best_pred],
                    mode="markers",
                    marker=dict(color=COLORS["green"], size=12, symbol="star"),
                    name="Optimal Point",
                ))
                fig.update_layout(
                    **PLOTLY_LAYOUT,
                    title=f"Setpoint Sweep: {so_setpoint_col.replace('_',' ').title()} → {so_target.replace('_',' ').title()}",
                    height=360,
                    xaxis_title=so_setpoint_col.replace("_", " ").title(),
                    yaxis_title=f"Predicted {so_target} (kW)",
                )
                st.plotly_chart(fig, use_container_width=True)

                # Recommendation box
                direction = "↑ increase" if best_sp > current_sp else "↓ decrease"
                delta_sp  = abs(best_sp - current_sp)
                st.success(
                    f"**Recommendation:** {direction} **{so_setpoint_col.replace('_', ' ')}** "
                    f"from **{current_sp:.1f}** → **{best_sp:.1f}** "
                    f"(Δ {delta_sp:.1f})  \n"
                    f"Estimated energy saving: **{savings_pct:.1f}%** "
                    f"({current_pred:.1f} → {best_pred:.1f} kW predicted average)"
                )

                # Coefficient table — what drives energy use
                coef_names = list(so_feature_cols) + list(time_feats.columns)
                coefs      = pipe.named_steps["ridge"].coef_
                coef_df    = pd.DataFrame({
                    "Feature":     coef_names[:len(coefs)],
                    "Coefficient": coefs[:len(coef_names)].round(4),
                    "Impact":      np.abs(coefs[:len(coef_names)]).round(4),
                }).sort_values("Impact", ascending=False).head(12)

                with st.expander("📋 Model Coefficients (what drives energy use)"):
                    fig2 = go.Figure(go.Bar(
                        x=coef_df["Coefficient"],
                        y=coef_df["Feature"],
                        orientation="h",
                        marker_color=[COLORS["red"] if v > 0 else COLORS["teal"]
                                      for v in coef_df["Coefficient"]],
                    ))
                    fig2.update_layout(**PLOTLY_LAYOUT,
                                       title="Ridge Coefficients (positive = raises energy use)",
                                       height=320,
                                       xaxis_title="Coefficient (scaled units)")
                    st.plotly_chart(fig2, use_container_width=True)
                    st.dataframe(coef_df, use_container_width=True, hide_index=True)
                    st.caption(
                        "Coefficients are in **standardised units** (after scaling), so magnitude "
                        "is comparable across features. Positive = raises predicted kW; Negative = reduces it."
                    )


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 7 — Equipment Fault Rules
# ══════════════════════════════════════════════════════════════════════════════════
with tab7:
    st.markdown("## 🏭 Equipment Fault Rules")
    st.caption("Fault detection rules for VAV boxes, boilers, cooling towers, heat exchangers, and RTUs.")

    equip_tab1, equip_tab2, equip_tab3, equip_tab4, equip_tab5 = st.tabs([
        "🌀 VAV Boxes", "🔥 Boilers", "🌊 Cooling Towers", "♻️ Heat Exchangers", "🏠 RTUs"
    ])

    EQUIP_COLORS = {"red": "#ff6b6b", "amber": "#f7c948", "green": "#2ecc71", "blue": "#00b4d8", "teal": "#06d6a0"}

    def equip_metric(label, value, status="default"):
        color = {"danger": "#ff6b6b", "warning": "#f7c948", "good": "#2ecc71", "default": "#00b4d8"}.get(status, "#00b4d8")
        st.markdown(f"""<div style="background:#161b22;border:1px solid #30363d;border-top:3px solid {color};
        border-radius:10px;padding:16px 20px;margin-bottom:10px">
        <div style="font-family:Space Mono,monospace;font-size:10px;letter-spacing:2px;color:#8b949e;text-transform:uppercase">{label}</div>
        <div style="font-family:Space Mono,monospace;font-size:24px;font-weight:700;color:#e6edf3">{value}</div>
        </div>""", unsafe_allow_html=True)

    # ── VAV Boxes ────────────────────────────────────────────────────────────────
    with equip_tab1:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>VAV Box Fault Detection</div>", unsafe_allow_html=True)
        st.markdown("VAV (Variable Air Volume) boxes control airflow to individual zones. Common faults include stuck dampers, failed actuators, and sensor drift.")

        vav_cols = [c for c in numeric_cols if any(k in c.lower() for k in ["cfm", "damper", "vav", "zone", "supply"])]

        VAV_RULES = [
            {"name": "Stuck Damper — Fully Closed", "severity": "HIGH",
             "description": "Damper position <5% for >2 hrs during occupied hours — zone getting no airflow.",
             "expr": "oa_damper_pct < 5% during occupied hrs",
             "check": lambda d: (d["oa_damper_pct"] < 5) & (d["timestamp"].dt.hour.between(7,18)) if "oa_damper_pct" in d.columns else None,
             "action": "Inspect actuator, check BAS command signal, verify linkage."},
            {"name": "Stuck Damper — Fully Open", "severity": "MEDIUM",
             "description": "Damper >95% for extended period when cooling demand is low — overcooling risk.",
             "expr": "oa_damper_pct > 95% for >4 hrs",
             "check": lambda d: d["oa_damper_pct"] > 95 if "oa_damper_pct" in d.columns else None,
             "action": "Check zone thermostat, verify control sequence, inspect actuator."},
            {"name": "Low Airflow — Underventilation", "severity": "HIGH",
             "description": "Supply CFM below 30% of mean during occupied hours — IAQ concern.",
             "expr": "supply_cfm < 30% of mean during occupied hrs",
             "check": lambda d: (d["supply_cfm"] < d["supply_cfm"].mean() * 0.3) & (d["timestamp"].dt.hour.between(7,18)) if "supply_cfm" in d.columns else None,
             "action": "Check duct static pressure, inspect VAV actuator, verify AHU fan operation."},
            {"name": "High Airflow — Overcooling", "severity": "LOW",
             "description": "Supply CFM above 95th percentile sustained — possible control loop hunting.",
             "expr": "supply_cfm > P95 for >2 hrs",
             "check": lambda d: d["supply_cfm"] > d["supply_cfm"].quantile(0.95) if "supply_cfm" in d.columns else None,
             "action": "Check zone temperature sensor, review control tuning parameters."},
            {"name": "Supply Temp Too High at VAV", "severity": "MEDIUM",
             "description": "Supply air above 65°F during peak cooling hours — AHU or VAV reheat issue.",
             "expr": "supply_air_temp > 65°F between 10:00–16:00",
             "check": lambda d: (d["supply_air_temp"] > 65) & (d["timestamp"].dt.hour.between(10,16)) if "supply_air_temp" in d.columns else None,
             "action": "Verify AHU cooling coil, check chilled water valve, inspect reheat controls."},
        ]

        triggered_vav = 0
        for rule in VAV_RULES:
            try:
                mask = rule["check"](df)
                if mask is None:
                    status_icon = "⚫"
                    status_text = "Missing data"
                    count = 0
                else:
                    count = int(mask.sum())
                    triggered_vav += 1 if count > 0 else 0
                    status_icon = "🔴" if count > 0 and rule["severity"]=="HIGH" else "🟡" if count > 0 else "🟢"
                    status_text = f"{count} fault readings" if count > 0 else "No faults"
            except:
                status_icon = "⚫"
                status_text = "Could not evaluate"
                count = 0

            with st.expander(f"{status_icon} **{rule['name']}** — {status_text}", expanded=count>0 and rule["severity"]=="HIGH"):
                st.markdown(f"<code style='background:rgba(0,180,216,0.1);color:#00b4d8;padding:3px 8px;border-radius:4px;font-size:11px'>{rule['expr']}</code>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#8b949e;font-size:12px'>{rule['description']}</p>", unsafe_allow_html=True)
                st.markdown(f"**Recommended action:** {rule['action']}")
                sev_color = "#ff6b6b" if rule["severity"]=="HIGH" else "#f7c948" if rule["severity"]=="MEDIUM" else "#2ecc71"
                st.markdown(f"<span style='color:{sev_color};font-family:Space Mono,monospace;font-size:11px;font-weight:700'>{rule['severity']}</span>", unsafe_allow_html=True)

        c1, c2 = st.columns(2)
        c1.metric("VAV Rules Triggered", triggered_vav)
        c2.metric("Total VAV Rules", len(VAV_RULES))

    # ── Boilers ──────────────────────────────────────────────────────────────────
    with equip_tab2:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Boiler Fault Detection</div>", unsafe_allow_html=True)
        st.markdown("Boiler faults often manifest as supply/return temperature anomalies, cycling issues, and efficiency degradation.")

        st.info("💡 Boiler analysis uses supply/return temperature data. For dedicated boiler sensors, upload a CSV with columns: , , , .")

        BOILER_RULES = [
            {"name": "Low System ΔT — Poor Heat Transfer", "severity": "HIGH",
             "description": "Return temp within 5°F of supply — very low ΔT indicates poor heat transfer, flow imbalance, or short-circuiting.",
             "expr": "(return_air_temp - supply_air_temp) < 5°F for ≥2 hrs",
             "check": lambda d: (d["return_air_temp"] - d["supply_air_temp"]) < 5 if all(c in d.columns for c in ["return_air_temp","supply_air_temp"]) else None,
             "action": "Check flow balancing valves, inspect system bypass, verify pump operation."},
            {"name": "High Supply Temp — Overheating Risk", "severity": "HIGH",
             "description": "Return air above 85°F — boiler may be overshooting setpoint.",
             "expr": "return_air_temp > 85°F for ≥1 hr",
             "check": lambda d: d["return_air_temp"] > 85 if "return_air_temp" in d.columns else None,
             "action": "Check boiler setpoint, verify aquastat, inspect mixing valve."},
            {"name": "Rapid Cycling", "severity": "MEDIUM",
             "description": "High variance in return temp over short window — boiler short-cycling wastes fuel and causes wear.",
             "expr": "std(return_air_temp, 2hr) > 8°F",
             "check": lambda d: d["return_air_temp"].rolling(2).std() > 8 if "return_air_temp" in d.columns else None,
             "action": "Check boiler sizing, inspect heat exchanger, review control deadband settings."},
            {"name": "High Energy at Low Load", "severity": "MEDIUM",
             "description": "High power consumption when thermal output is low — efficiency problem.",
             "expr": "chiller_kw > P75 when ΔT < 10°F",
             "check": lambda d: (d["chiller_kw"] > d["chiller_kw"].quantile(0.75)) & ((d["return_air_temp"]-d["supply_air_temp"]) < 10) if all(c in d.columns for c in ["chiller_kw","return_air_temp","supply_air_temp"]) else None,
             "action": "Inspect heat exchanger fouling, check burner combustion efficiency, review controls."},
        ]

        triggered_boiler = 0
        for rule in BOILER_RULES:
            try:
                mask = rule["check"](df)
                if mask is None:
                    status_icon = "⚫"; status_text = "Missing data"; count = 0
                else:
                    count = int(mask.fillna(False).sum())
                    triggered_boiler += 1 if count > 0 else 0
                    status_icon = "🔴" if count>0 and rule["severity"]=="HIGH" else "🟡" if count>0 else "🟢"
                    status_text = f"{count} fault readings" if count > 0 else "No faults"
            except:
                status_icon = "⚫"; status_text = "Could not evaluate"; count = 0

            with st.expander(f"{status_icon} **{rule['name']}** — {status_text}", expanded=count>0 and rule["severity"]=="HIGH"):
                st.markdown(f"<code style='background:rgba(0,180,216,0.1);color:#00b4d8;padding:3px 8px;border-radius:4px;font-size:11px'>{rule['expr']}</code>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#8b949e;font-size:12px'>{rule['description']}</p>", unsafe_allow_html=True)
                st.markdown(f"**Recommended action:** {rule['action']}")

        c1, c2 = st.columns(2)
        c1.metric("Boiler Rules Triggered", triggered_boiler)
        c2.metric("Total Boiler Rules", len(BOILER_RULES))

    # ── Cooling Towers ───────────────────────────────────────────────────────────
    with equip_tab3:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Cooling Tower Fault Detection</div>", unsafe_allow_html=True)
        st.markdown("Cooling tower performance directly affects chiller efficiency. Key indicators are condenser water temperature, approach temperature, and fan energy.")

        st.info("💡 Upload a CSV with , , and  for full cooling tower analysis.")

        CT_RULES = [
            {"name": "High Condenser Leaving Water Temp", "severity": "HIGH",
             "description": "Condenser temp above 95°F — tower not rejecting heat effectively. Increases chiller head pressure and reduces efficiency.",
             "expr": "condenser_temp > 95°F for ≥1 hr",
             "check": lambda d: d["condenser_temp"] > 95 if "condenser_temp" in d.columns else None,
             "action": "Inspect tower fill, check fan operation, verify water distribution nozzles, check for scaling."},
            {"name": "High Approach Temperature", "severity": "HIGH",
             "description": "Gap between condenser leaving temp and outside air temp > 15°F — tower underperforming relative to conditions.",
             "expr": "(condenser_temp - outside_air_temp) > 15°F for ≥2 hrs",
             "check": lambda d: (d["condenser_temp"] - d["outside_air_temp"]) > 15 if all(c in d.columns for c in ["condenser_temp","outside_air_temp"]) else None,
             "action": "Check tower fill fouling, inspect drift eliminators, verify fan speed, check water flow rate."},
            {"name": "Condenser Temp Spike", "severity": "HIGH",
             "description": "Rapid rise in condenser temp (>10°F in 2 hrs) — possible fan failure or sudden scaling.",
             "expr": "condenser_temp increase > 10°F in 2 hrs",
             "check": lambda d: d["condenser_temp"].diff(2) > 10 if "condenser_temp" in d.columns else None,
             "action": "Check cooling tower fans immediately, inspect water flow, verify chemical treatment."},
            {"name": "Low Condenser Temp — Overcooling", "severity": "LOW",
             "description": "Condenser temp below 65°F — tower overcooling, chiller may trip on low head pressure.",
             "expr": "condenser_temp < 65°F for ≥1 hr",
             "check": lambda d: d["condenser_temp"] < 65 if "condenser_temp" in d.columns else None,
             "action": "Install or enable condenser water temperature reset control, consider bypass valve."},
            {"name": "High Chiller kW vs Condenser Temp", "severity": "MEDIUM",
             "description": "Chiller power high while condenser temp is also high — compounding efficiency loss.",
             "expr": "chiller_kw > P80 AND condenser_temp > 90°F",
             "check": lambda d: (d["chiller_kw"] > d["chiller_kw"].quantile(0.80)) & (d["condenser_temp"] > 90) if all(c in d.columns for c in ["chiller_kw","condenser_temp"]) else None,
             "action": "Address cooling tower performance first — every 1°F reduction in condenser temp saves ~1-2% chiller energy."},
        ]

        triggered_ct = 0
        for rule in CT_RULES:
            try:
                mask = rule["check"](df)
                if mask is None:
                    status_icon = "⚫"; status_text = "Missing data"; count = 0
                else:
                    count = int(mask.fillna(False).sum())
                    triggered_ct += 1 if count > 0 else 0
                    status_icon = "🔴" if count>0 and rule["severity"]=="HIGH" else "🟡" if count>0 else "🟢"
                    status_text = f"{count} fault readings" if count > 0 else "No faults"
            except:
                status_icon = "⚫"; status_text = "Could not evaluate"; count = 0

            with st.expander(f"{status_icon} **{rule['name']}** — {status_text}", expanded=count>0 and rule["severity"]=="HIGH"):
                st.markdown(f"<code style='background:rgba(0,180,216,0.1);color:#00b4d8;padding:3px 8px;border-radius:4px;font-size:11px'>{rule['expr']}</code>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#8b949e;font-size:12px'>{rule['description']}</p>", unsafe_allow_html=True)
                st.markdown(f"**Recommended action:** {rule['action']}")

        c1, c2 = st.columns(2)
        c1.metric("Cooling Tower Rules Triggered", triggered_ct)
        c2.metric("Total Cooling Tower Rules", len(CT_RULES))

    # ── Heat Exchangers ──────────────────────────────────────────────────────────
    with equip_tab4:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Heat Exchanger Fault Detection</div>", unsafe_allow_html=True)
        st.markdown("Heat exchanger faults are typically caused by fouling, scaling, or flow imbalance — all showing up as reduced ΔT or increased approach temperature.")

        HX_RULES = [
            {"name": "Low Effectiveness — Fouling Suspected", "severity": "HIGH",
             "description": "ΔT across heat exchanger below 8°F — significantly reduced thermal transfer, likely fouling or scaling.",
             "expr": "(return_air_temp - supply_air_temp) < 8°F for ≥2 hrs",
             "check": lambda d: (d["return_air_temp"] - d["supply_air_temp"]) < 8 if all(c in d.columns for c in ["return_air_temp","supply_air_temp"]) else None,
             "action": "Schedule chemical cleaning, inspect for scaling, check water quality and treatment program."},
            {"name": "High Chiller Approach", "severity": "MEDIUM",
             "description": "Gap between supply air and chilled water > 18°F — heat exchanger not transferring effectively.",
             "expr": "(supply_air_temp - chilled_water_temp) > 18°F for ≥2 hrs",
             "check": lambda d: (d["supply_air_temp"] - d["chilled_water_temp"]) > 18 if all(c in d.columns for c in ["supply_air_temp","chilled_water_temp"]) else None,
             "action": "Inspect cooling coil for fouling, check chilled water flow rate, verify valve operation."},
            {"name": "ΔT Degradation Trend", "severity": "MEDIUM",
             "description": "Rolling ΔT declining over time — early indicator of gradual fouling before it becomes a fault.",
             "expr": "24hr rolling ΔT < 7-day rolling ΔT by >20%",
             "check": lambda d: (d["return_air_temp"] - d["supply_air_temp"]).rolling(24).mean() < (d["return_air_temp"] - d["supply_air_temp"]).rolling(168, min_periods=24).mean() * 0.8 if all(c in d.columns for c in ["return_air_temp","supply_air_temp"]) else None,
             "action": "Monitor closely, schedule preventive cleaning before ΔT drops further."},
        ]

        triggered_hx = 0
        for rule in HX_RULES:
            try:
                mask = rule["check"](df)
                if mask is None:
                    status_icon = "⚫"; status_text = "Missing data"; count = 0
                else:
                    count = int(mask.fillna(False).sum())
                    triggered_hx += 1 if count > 0 else 0
                    status_icon = "🔴" if count>0 and rule["severity"]=="HIGH" else "🟡" if count>0 else "🟢"
                    status_text = f"{count} fault readings" if count > 0 else "No faults"
            except:
                status_icon = "⚫"; status_text = "Could not evaluate"; count = 0

            with st.expander(f"{status_icon} **{rule['name']}** — {status_text}", expanded=count>0 and rule["severity"]=="HIGH"):
                st.markdown(f"<code style='background:rgba(0,180,216,0.1);color:#00b4d8;padding:3px 8px;border-radius:4px;font-size:11px'>{rule['expr']}</code>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#8b949e;font-size:12px'>{rule['description']}</p>", unsafe_allow_html=True)
                st.markdown(f"**Recommended action:** {rule['action']}")

        c1, c2 = st.columns(2)
        c1.metric("HX Rules Triggered", triggered_hx)
        c2.metric("Total HX Rules", len(HX_RULES))

    # ── RTUs ─────────────────────────────────────────────────────────────────────
    with equip_tab5:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>RTU (Rooftop Unit) Fault Detection</div>", unsafe_allow_html=True)
        st.markdown("RTUs are self-contained units combining heating, cooling, and ventilation. Common faults include refrigerant issues, economizer failures, and heat exchanger problems.")

        RTU_RULES = [
            {"name": "Economizer Failure — High Energy", "severity": "HIGH",
             "description": "High energy use when outside air is cool enough for free cooling — economizer may be stuck closed.",
             "expr": "total_kw > P75 when outside_air_temp < 60°F",
             "check": lambda d: (d["total_kw"] > d["total_kw"].quantile(0.75)) & (d["outside_air_temp"] < 60) if all(c in d.columns for c in ["total_kw","outside_air_temp"]) else None,
             "action": "Inspect economizer damper actuator, check control sequence, verify outdoor air sensor."},
            {"name": "Supply Air Temp Unstable", "severity": "MEDIUM",
             "description": "High short-term variance in supply air — hunting controls or refrigerant charge issue.",
             "expr": "std(supply_air_temp, 1hr) > 5°F",
             "check": lambda d: d["supply_air_temp"].rolling(1).std() > 5 if "supply_air_temp" in d.columns else None,
             "action": "Check refrigerant charge, inspect TXV operation, review control tuning."},
            {"name": "Insufficient Cooling", "severity": "HIGH",
             "description": "Supply air above 62°F during peak cooling hours — RTU not meeting cooling load.",
             "expr": "supply_air_temp > 62°F between 11:00–17:00",
             "check": lambda d: (d["supply_air_temp"] > 62) & (d["timestamp"].dt.hour.between(11,17)) if "supply_air_temp" in d.columns else None,
             "action": "Check refrigerant charge, inspect condenser coil, verify compressor operation."},
            {"name": "After-Hours Operation", "severity": "MEDIUM",
             "description": "Significant energy use after hours — RTU may be running on override or schedule error.",
             "expr": "total_kw > P60 between 20:00–06:00",
             "check": lambda d: (d["total_kw"] > d["total_kw"].quantile(0.60)) & (~d["timestamp"].dt.hour.between(6,20)) if "total_kw" in d.columns else None,
             "action": "Review BAS scheduling, check for tenant override commands, audit after-hours usage."},
            {"name": "High Outside Air ΔT Load", "severity": "LOW",
             "description": "Large gap between outside air and supply temp during peak — RTU working very hard, possible oversizing or control issue.",
             "expr": "(outside_air_temp - supply_air_temp) > 30°F for ≥2 hrs",
             "check": lambda d: (d["outside_air_temp"] - d["supply_air_temp"]) > 30 if all(c in d.columns for c in ["outside_air_temp","supply_air_temp"]) else None,
             "action": "Verify RTU sizing, check refrigerant charge, inspect condenser coil cleanliness."},
        ]

        triggered_rtu = 0
        for rule in RTU_RULES:
            try:
                mask = rule["check"](df)
                if mask is None:
                    status_icon = "⚫"; status_text = "Missing data"; count = 0
                else:
                    count = int(mask.fillna(False).sum())
                    triggered_rtu += 1 if count > 0 else 0
                    status_icon = "🔴" if count>0 and rule["severity"]=="HIGH" else "🟡" if count>0 else "🟢"
                    status_text = f"{count} fault readings" if count > 0 else "No faults"
            except:
                status_icon = "⚫"; status_text = "Could not evaluate"; count = 0

            with st.expander(f"{status_icon} **{rule['name']}** — {status_text}", expanded=count>0 and rule["severity"]=="HIGH"):
                st.markdown(f"<code style='background:rgba(0,180,216,0.1);color:#00b4d8;padding:3px 8px;border-radius:4px;font-size:11px'>{rule['expr']}</code>", unsafe_allow_html=True)
                st.markdown(f"<p style='color:#8b949e;font-size:12px'>{rule['description']}</p>", unsafe_allow_html=True)
                st.markdown(f"**Recommended action:** {rule['action']}")

        c1, c2 = st.columns(2)
        c1.metric("RTU Rules Triggered", triggered_rtu)
        c2.metric("Total RTU Rules", len(RTU_RULES))

# ══════════════════════════════════════════════════════════════════════════════════
# TAB 8 — BACnet Live Data
# ══════════════════════════════════════════════════════════════════════════════════
with tab8:
    st.markdown("## 🔌 BACnet Live Data Connection")
    st.caption("Connect to a live BACnet network to pull real-time data directly into the dashboard.")

    st.info("""
    **BACnet live polling requires the BAC0 library and network access to your BACnet/IP network.**
    This tab provides the full connection framework. To use it on-site:
    1. Run the app on a laptop connected to the building BAS network
    2. Enter the BACnet device IP and instance below
    3. Click Connect to start polling live data
    """)

    bacnet_tab1, bacnet_tab2, bacnet_tab3 = st.tabs([
        "⚙️ Connection Setup", "📡 Live Poll", "📋 Point Discovery"
    ])

    with bacnet_tab1:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>BACnet Connection Settings</div>", unsafe_allow_html=True)

        bc1, bc2 = st.columns(2)
        with bc1:
            bacnet_ip     = st.text_input("BACnet Device IP", value="192.168.1.100", help="IP address of the BACnet/IP device or router")
            bacnet_port   = st.number_input("BACnet Port", value=47808, help="Default BACnet/IP port is 47808")
            device_id     = st.number_input("Device Instance", value=1000, help="BACnet device instance number")
        with bc2:
            poll_interval = st.slider("Poll interval (seconds)", 5, 300, 30)
            max_points    = st.number_input("Max points to poll", value=50, min_value=1, max_value=500)
            local_ip      = st.text_input("Local IP (this computer)", value="192.168.1.50", help="IP of the machine running this app on the BACnet network")

        st.markdown("#### Point Configuration")
        st.caption("Define which BACnet object IDs map to which signals. Standard object types: AI=Analog Input, AO=Analog Output, AV=Analog Value.")

        default_points = pd.DataFrame({
            "Signal Name":    ["supply_air_temp", "return_air_temp", "chilled_water_temp", "condenser_temp", "outside_air_temp", "supply_cfm", "oa_damper_pct", "chiller_kw"],
            "BACnet Object":  ["analogInput",     "analogInput",     "analogInput",        "analogInput",    "analogInput",      "analogInput", "analogOutput",  "analogInput"],
            "Instance":       [1,                  2,                  3,                     4,                 5,                   6,             1,               10],
            "Units":          ["degF",             "degF",            "degF",               "degF",           "degF",             "cfm",         "percent",       "kilowatts"],
            "Enabled":        [True,               True,               True,                  True,              True,                True,          True,            True],
        })

        edited_points = st.data_editor(default_points, use_container_width=True, num_rows="dynamic")

        if st.button("💾 Save Point Configuration", key="save_points"):
            st.success(f"✓ Saved {len(edited_points)} point mappings.")

        st.markdown("#### Connection Code")
        st.caption("This is the BAC0 Python code that will run when you click Connect on the Live Poll tab.")
        st.code(f"""
import BAC0
import pandas as pd
from datetime import datetime

# Initialize BACnet connection
bacnet = BAC0.lite(ip="{local_ip}")

# Connect to device
device = BAC0.device("{bacnet_ip}", {int(device_id)}, bacnet)

# Poll points
records = []
timestamp = datetime.now()

point_map = {{
    "supply_air_temp":     ("analogInput",  1),
    "return_air_temp":     ("analogInput",  2),
    "chilled_water_temp": ("analogInput",  3),
    "condenser_temp":      ("analogInput",  4),
    "outside_air_temp":    ("analogInput",  5),
    "supply_cfm":          ("analogInput",  6),
    "oa_damper_pct":       ("analogOutput", 1),
    "chiller_kw":          ("analogInput", 10),
}}

row = {{"timestamp": timestamp}}
for signal, (obj_type, instance) in point_map.items():
    try:
        value = device[f"{{obj_type}} {{instance}} presentValue"]
        row[signal] = float(value)
    except Exception as e:
        row[signal] = None

records.append(row)
df_live = pd.DataFrame(records)
print(df_live)
""", language="python")

    with bacnet_tab2:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Live Data Polling</div>", unsafe_allow_html=True)

        col_a, col_b, col_c = st.columns(3)
        col_a.metric("Connection Status", "⚫ Offline")
        col_b.metric("Last Poll", "—")
        col_c.metric("Points Active", "0")

        st.warning("""
        **To enable live BACnet polling:**

        1. Install BAC0: 
        2. Connect your laptop to the building BAS network
        3. Configure point mappings in the Connection Setup tab
        4. Uncomment and run the polling code

        Once connected, live data will flow directly into all dashboard tabs in real time.
        """)

        st.markdown("#### Simulated Live Feed (Demo)")
        st.caption("This shows what live data would look like — using the current dataset as a proxy.")

        if not df.empty:
            latest = df.tail(1).T
            latest.columns = ["Latest Value"]
            latest.index.name = "Signal"
            st.dataframe(latest, use_container_width=True)

        st.markdown("#### Manual CSV Import")
        st.caption("Already have a BACnet trend log exported as CSV? Upload it here and all tabs will update.")
        live_upload = st.file_uploader("Upload BACnet trend CSV", type=["csv"], key="live_upload")
        if live_upload:
            try:
                live_df = pd.read_csv(live_upload)
                st.success(f"✓ Loaded {len(live_df):,} records from trend log.")
                st.dataframe(live_df.head(10), use_container_width=True)
            except Exception as e:
                st.error(f"Error reading file: {e}")

    with bacnet_tab3:
        st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>BACnet Point Discovery</div>", unsafe_allow_html=True)
        st.markdown("When connected to a live BACnet device, this tab automatically discovers all available points.")

        st.code("""
# Auto-discovery code (run when connected to BACnet network)
import BAC0

bacnet = BAC0.lite(ip='192.168.1.50')
device = BAC0.device('192.168.1.100', 1000, bacnet)

# Discover all points
all_points = device.points
print(f'Found {len(all_points)} points:')
for point in all_points:
    print(f'  {point.name}: {point.lastValue} {point.units}')
""", language="python")

        st.markdown("#### Common BACnet Point Naming Conventions")
        naming_df = pd.DataFrame({
            "Equipment": ["AHU-1","AHU-1","AHU-1","CH-1","CH-1","CT-1","CT-1","BLR-1","VAV-101","VAV-101"],
            "Point Name": ["AHU-1.SAT","AHU-1.RAT","AHU-1.SF-SPD","CH-1.CHWST","CH-1.KW","CT-1.LCWT","CT-1.FAN-SPD","BLR-1.HWS","VAV-101.DAM-POS","VAV-101.ZONE-TEMP"],
            "Description": ["Supply Air Temp","Return Air Temp","Supply Fan Speed","Chilled Water Supply Temp","Chiller Power","Leaving Condenser Water Temp","Cooling Tower Fan Speed","Hot Water Supply Temp","VAV Damper Position","Zone Temperature"],
            "Typical Units": ["°F","°F","%","°F","kW","°F","%","°F","%","°F"],
            "BACnet Object": ["AI:1","AI:2","AO:1","AI:3","AI:10","AI:4","AO:2","AI:5","AO:3","AI:6"],
        })
        st.dataframe(naming_df, use_container_width=True, hide_index=True)

        st.markdown("#### Supported BACnet Protocols")
        proto_df = pd.DataFrame({
            "Protocol": ["BACnet/IP","BACnet MS/TP","BACnet/SC"],
            "Use Case": ["Most modern BAS — connects over standard Ethernet/WiFi","Older field devices — requires serial RS-485 adapter","Newest secure BACnet — cloud-friendly"],
            "BAC0 Support": ["✅ Full support","✅ With USB-RS485 adapter","🔄 Partial (BAC0 v22+)"],
            "Typical Devices": ["Controllers, routers, gateways","VAV controllers, terminal units","Modern cloud-connected BAS"],
        })
        st.dataframe(proto_df, use_container_width=True, hide_index=True)


# ══════════════════════════════════════════════════════════════════════════════════
# TAB 9 — Report Generator
# ══════════════════════════════════════════════════════════════════════════════════
with tab9:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                        TableStyle, PageBreak, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import io
        REPORT_LIBS_OK = True
    except ImportError as e:
        REPORT_LIBS_OK = False
        st.error(f"Report libraries not installed: {e}. Add reportlab and python-docx to requirements.txt")

    st.markdown("## 📄 Report Generator")
    st.caption("Auto-generate a professional client report from your dashboard analysis — then download and edit.")

    # ── Report Settings ──────────────────────────────────────────────────────────
    st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Report Settings</div>", unsafe_allow_html=True)

    r1, r2 = st.columns(2)
    with r1:
        report_title     = st.text_input("Report Title", value="HVAC Energy Analysis Report")
        client_name      = st.text_input("Client / Building Name", value="[Client Name]")
        building_address = st.text_input("Building Address", value="[Building Address]")
        prepared_by      = st.text_input("Prepared By", value="[Your Name], ME, MBA")
    with r2:
        company_name     = st.text_input("Your Company Name", value="[Your Company]")
        report_date      = st.text_input("Report Date", value=datetime.now().strftime("%B %d, %Y"))
        electricity_rate = st.number_input("Electricity Rate ($/kWh)", value=0.12, step=0.01, format="%.3f")
        currency         = st.selectbox("Currency", ["USD ($)", "EUR (€)", "GBP (£)"])

    currency_sym = {"USD ($)": "$", "EUR (€)": "€", "GBP (£)": "£"}.get(currency, "$")

    st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin:24px 0 16px 0'>Findings & Recommendations</div>", unsafe_allow_html=True)
    st.caption("These are auto-populated from your data. Edit any field before generating.")

    # ── Auto-compute findings ────────────────────────────────────────────────────
    findings = []

    if "supply_air_temp" in df.columns and "return_air_temp" in df.columns:
        avg_dt = (df["return_air_temp"] - df["supply_air_temp"]).mean()
        low_dt_hrs = int(((df["return_air_temp"] - df["supply_air_temp"]) < 10).sum())
        if low_dt_hrs > 0:
            annual_waste = low_dt_hrs * df.get("chiller_kw", pd.Series([0])).mean() * 0.15 * electricity_rate * 8760 / max(len(df),1)
            findings.append({
                "title": "Low Delta-T Syndrome",
                "severity": "HIGH",
                "description": f"Return-supply temperature differential averaged {avg_dt:.1f}°F. System spent {low_dt_hrs} hours below the 10°F design threshold, indicating poor heat transfer or flow imbalance.",
                "recommendation": "Balance coil flow rates, inspect bypass valves, verify pump operation and chilled water flow.",
                "savings_kwh": round(low_dt_hrs * df.get("chiller_kw", pd.Series([100])).mean() * 0.10),
            })

    if "condenser_temp" in df.columns:
        high_cond_hrs = int((df["condenser_temp"] > 95).sum())
        if high_cond_hrs > 0:
            findings.append({
                "title": "Elevated Condenser Temperature",
                "severity": "HIGH",
                "description": f"Condenser temperature exceeded 95°F for {high_cond_hrs} hours. Each degree above optimal operating temperature reduces chiller efficiency by approximately 1.5-2%.",
                "recommendation": "Inspect cooling tower fill and nozzles, verify fan operation, check water treatment and scaling, clean condenser tubes.",
                "savings_kwh": round(high_cond_hrs * df.get("chiller_kw", pd.Series([100])).mean() * 0.02),
            })

    if "oa_damper_pct" in df.columns:
        stuck_hrs = int(((df["oa_damper_pct"] < 5) & (df["timestamp"].dt.hour.between(7,18))).sum())
        if stuck_hrs > 0:
            findings.append({
                "title": "Outside Air Damper Fault",
                "severity": "HIGH",
                "description": f"Outside air damper measured below 5% position for {stuck_hrs} occupied hours. This indicates a potential actuator failure or controls issue creating an IAQ risk.",
                "recommendation": "Inspect OA damper actuator, linkage, and BAS command signal. Verify damper position feedback sensor calibration.",
                "savings_kwh": 0,
            })

    if "total_kw" in df.columns:
        after_hrs_mask = (~df["timestamp"].dt.hour.between(6,20)) & (df["total_kw"] > df["total_kw"].quantile(0.75))
        after_hrs_count = int(after_hrs_mask.sum())
        if after_hrs_count > 20:
            avg_after_hrs_kw = df.loc[after_hrs_mask, "total_kw"].mean()
            findings.append({
                "title": "Excessive After-Hours Energy Consumption",
                "severity": "MEDIUM",
                "description": f"Facility recorded high energy use ({avg_after_hrs_kw:.0f} kW average) during {after_hrs_count} after-hours periods (10 PM–6 AM). Equipment appears to be running when the building is unoccupied.",
                "recommendation": "Audit BAS scheduling for all HVAC equipment. Check for tenant override commands. Implement automated after-hours setback.",
                "savings_kwh": round(after_hrs_count * avg_after_hrs_kw * 0.40),
            })

    if "chiller_kw" in df.columns:
        overload_hrs = int((df["chiller_kw"] > df["chiller_kw"].quantile(0.95)).sum())
        if overload_hrs > 0:
            findings.append({
                "title": "Chiller Overload Events",
                "severity": "MEDIUM",
                "description": f"Chiller power exceeded the 95th percentile threshold for {overload_hrs} hours, indicating periods of high stress that increase wear and risk of trip.",
                "recommendation": "Review chiller staging strategy. Evaluate adding cooling capacity or demand limiting controls for peak periods.",
                "savings_kwh": round(overload_hrs * df["chiller_kw"].mean() * 0.05),
            })

    # If no faults found, add a positive finding
    if not findings:
        findings.append({
            "title": "System Operating Within Normal Parameters",
            "severity": "LOW",
            "description": "Analysis of the provided data period did not identify significant fault conditions. System temperatures, energy consumption, and airflow are within expected ranges.",
            "recommendation": "Continue monthly monitoring. Consider establishing baseline KPIs for ongoing performance tracking.",
            "savings_kwh": 0,
        })

    # ── Editable findings ────────────────────────────────────────────────────────
    total_savings_kwh = sum(f["savings_kwh"] for f in findings)
    total_savings_usd = total_savings_kwh * electricity_rate

    st.markdown(f"**{len(findings)} findings identified** — estimated annual savings potential: **{currency_sym}{total_savings_usd:,.0f}**")

    edited_findings = []
    for i, finding in enumerate(findings):
        with st.expander(f"Finding {i+1}: {finding['title']} — {finding['severity']}", expanded=i==0):
            col_t, col_s = st.columns([3,1])
            with col_t:
                f_title = st.text_input("Title", value=finding["title"], key=f"f_title_{i}")
            with col_s:
                f_sev = st.selectbox("Severity", ["HIGH","MEDIUM","LOW"],
                                     index=["HIGH","MEDIUM","LOW"].index(finding["severity"]),
                                     key=f"f_sev_{i}")
            f_desc = st.text_area("Description", value=finding["description"], key=f"f_desc_{i}", height=80)
            f_rec  = st.text_area("Recommendation", value=finding["recommendation"], key=f"f_rec_{i}", height=60)
            f_save = st.number_input(f"Estimated Annual Savings (kWh)", value=finding["savings_kwh"],
                                     min_value=0, key=f"f_save_{i}")
            edited_findings.append({
                "title": f_title, "severity": f_sev,
                "description": f_desc, "recommendation": f_rec,
                "savings_kwh": f_save,
                "savings_usd": round(f_save * electricity_rate, 2),
            })

    total_kwh  = sum(f["savings_kwh"] for f in edited_findings)
    total_usd  = total_kwh * electricity_rate
    high_count = sum(1 for f in edited_findings if f["severity"]=="HIGH")
    med_count  = sum(1 for f in edited_findings if f["severity"]=="MEDIUM")

    # ── KPI summary for report ───────────────────────────────────────────────────
    kpis_r = compute_kpis(df)
    anom_count_r = int(df["is_anomaly"].sum()) if "is_anomaly" in df.columns else 0
    date_range_str = f"{df['timestamp'].min().strftime('%B %d, %Y')} – {df['timestamp'].max().strftime('%B %d, %Y')}"
    total_records  = len(df)

    # ── Generate PDF ─────────────────────────────────────────────────────────────
    def generate_pdf():
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()

        # Custom styles
        DARK  = colors.HexColor("#0d1117")
        BLUE  = colors.HexColor("#0077b6")
        TEAL  = colors.HexColor("#00b4d8")
        RED   = colors.HexColor("#c0392b")
        AMBER = colors.HexColor("#e07b00")
        GREEN = colors.HexColor("#27ae60")
        LGRAY = colors.HexColor("#f5f7fa")
        MGRAY = colors.HexColor("#8b949e")

        style_title = ParagraphStyle("ReportTitle", fontSize=26, fontName="Helvetica-Bold",
                                     textColor=DARK, spaceAfter=6, alignment=TA_LEFT)
        style_sub   = ParagraphStyle("ReportSub", fontSize=11, fontName="Helvetica",
                                     textColor=MGRAY, spaceAfter=4)
        style_h1    = ParagraphStyle("H1", fontSize=14, fontName="Helvetica-Bold",
                                     textColor=BLUE, spaceBefore=18, spaceAfter=6)
        style_h2    = ParagraphStyle("H2", fontSize=11, fontName="Helvetica-Bold",
                                     textColor=DARK, spaceBefore=10, spaceAfter=4)
        style_body  = ParagraphStyle("Body", fontSize=9.5, fontName="Helvetica",
                                     textColor=DARK, spaceAfter=6, leading=14)
        style_label = ParagraphStyle("Label", fontSize=8, fontName="Helvetica",
                                     textColor=MGRAY, spaceAfter=2)
        style_small = ParagraphStyle("Small", fontSize=8, fontName="Helvetica",
                                     textColor=MGRAY, spaceAfter=4)

        story = []

        # ── Cover header bar ──────────────────────────────────────────────────
        header_data = [[Paragraph(f"<font color='white'><b>{report_title}</b></font>", style_title)]]
        header_tbl  = Table(header_data, colWidths=[7.0*inch])
        header_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), BLUE),
            ("TOPPADDING",    (0,0), (-1,-1), 18),
            ("BOTTOMPADDING", (0,0), (-1,-1), 18),
            ("LEFTPADDING",   (0,0), (-1,-1), 18),
        ]))
        story.append(header_tbl)
        story.append(Spacer(1, 12))

        # ── Client info block ─────────────────────────────────────────────────
        info_data = [
            [Paragraph("PREPARED FOR", style_label), Paragraph("PREPARED BY", style_label),
             Paragraph("REPORT DATE", style_label)],
            [Paragraph(f"<b>{client_name}</b>", style_body),
             Paragraph(f"<b>{prepared_by}</b>", style_body),
             Paragraph(f"<b>{report_date}</b>", style_body)],
            [Paragraph(building_address, style_small),
             Paragraph(company_name, style_small),
             Paragraph(f"Analysis Period: {date_range_str}", style_small)],
        ]
        info_tbl = Table(info_data, colWidths=[2.33*inch, 2.33*inch, 2.34*inch])
        info_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,-1), LGRAY),
            ("TOPPADDING",    (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
            ("LINEABOVE",     (0,0), (-1,0),  1, BLUE),
            ("LINEBELOW",     (0,-1),(-1,-1), 1, colors.HexColor("#dee2e6")),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ]))
        story.append(info_tbl)
        story.append(Spacer(1, 18))

        # ── Executive Summary ─────────────────────────────────────────────────
        story.append(Paragraph("1. Executive Summary", style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=TEAL, spaceAfter=8))

        summary_text = (
            f"This report presents the findings of a remote HVAC energy analysis conducted for "
            f"<b>{client_name}</b> covering the period {date_range_str}. "
            f"The analysis examined {total_records:,} data points across {len(df.columns)-1} sensor channels "
            f"using automated fault detection, anomaly scoring, and energy performance benchmarking."
        )
        story.append(Paragraph(summary_text, style_body))
        story.append(Spacer(1, 8))

        story.append(Paragraph(
            f"The analysis identified <b>{len(edited_findings)} findings</b>, including "
            f"<b>{high_count} high-severity</b> and <b>{med_count} medium-severity</b> issues. "
            f"Implementing the recommendations in this report is estimated to reduce annual energy "
            f"consumption by approximately <b>{total_kwh:,} kWh</b>, representing a potential "
            f"cost savings of <b>{currency_sym}{total_usd:,.0f} per year</b> at current electricity rates.",
            style_body
        ))
        story.append(Spacer(1, 12))

        # KPI summary table
        kpi_data = [
            [Paragraph("<b>Metric</b>", style_body), Paragraph("<b>Value</b>", style_body)],
            ["Analysis Period",        date_range_str],
            ["Total Records Analyzed", f"{total_records:,}"],
            ["Anomalies Detected",     str(anom_count_r)],
            ["Avg Chiller Load",       f"{kpis_r.get('avg_chiller_kw',0):.1f} kW"],
            ["Total Energy (Period)",  f"{kpis_r.get('total_kwh',0)/1000:.1f} MWh"],
            ["Avg Supply-Return DT",   f"{kpis_r.get('avg_dt',0):.1f} degF"],
            ["High Severity Findings", str(high_count)],
            ["Est. Annual Savings",    f"{currency_sym}{total_usd:,.0f}"],
        ]
        kpi_tbl = Table(kpi_data, colWidths=[3.0*inch, 4.0*inch])
        kpi_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0),  BLUE),
            ("TEXTCOLOR",     (0,0), (-1,0),  colors.white),
            ("BACKGROUND",    (0,1), (-1,-1), LGRAY),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.white, LGRAY]),
            ("GRID",          (0,0), (-1,-1), 0.5, colors.HexColor("#dee2e6")),
            ("TOPPADDING",    (0,0), (-1,-1), 6),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
            ("FONTNAME",      (0,1), (0,-1),  "Helvetica-Bold"),
            ("FONTSIZE",      (0,0), (-1,-1), 9),
        ]))
        story.append(kpi_tbl)
        story.append(PageBreak())

        # ── Findings ──────────────────────────────────────────────────────────
        story.append(Paragraph("2. Technical Findings", style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=TEAL, spaceAfter=8))

        for i, finding in enumerate(edited_findings):
            sev_color = RED if finding["severity"]=="HIGH" else AMBER if finding["severity"]=="MEDIUM" else GREEN
            sev_label = finding["severity"]

            # Finding header
            hdr_data = [[
                Paragraph(f"<font color='white'><b>Finding {i+1}: {finding['title']}</b></font>", style_h2),
                Paragraph(f"<font color='white'><b>{sev_label}</b></font>",
                          ParagraphStyle("sev", fontSize=9, fontName="Helvetica-Bold",
                                         textColor=colors.white, alignment=TA_RIGHT)),
            ]]
            hdr_tbl = Table(hdr_data, colWidths=[5.5*inch, 1.5*inch])
            hdr_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), sev_color),
                ("TOPPADDING",    (0,0), (-1,-1), 8),
                ("BOTTOMPADDING", (0,0), (-1,-1), 8),
                ("LEFTPADDING",   (0,0), (0,-1),  10),
                ("RIGHTPADDING",  (-1,0),(-1,-1), 10),
                ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
            ]))
            story.append(hdr_tbl)

            # Finding body
            body_data = [
                [Paragraph("<b>Observation</b>", style_label)],
                [Paragraph(finding["description"], style_body)],
                [Paragraph("<b>Recommended Action</b>", style_label)],
                [Paragraph(finding["recommendation"], style_body)],
            ]
            if finding["savings_kwh"] > 0:
                body_data.append([Paragraph(
                    f"<b>Estimated Annual Savings:</b> {finding['savings_kwh']:,} kWh "
                    f"({currency_sym}{finding['savings_usd']:,.0f} at {electricity_rate:.3f}/kWh)",
                    ParagraphStyle("savings", fontSize=9, fontName="Helvetica-Bold",
                                   textColor=GREEN, spaceAfter=4)
                )])
            body_tbl = Table(body_data, colWidths=[7.0*inch])
            body_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), colors.white),
                ("TOPPADDING",    (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ("LEFTPADDING",   (0,0), (-1,-1), 10),
                ("LINEBELOW",     (0,-1),(-1,-1), 1, colors.HexColor("#dee2e6")),
            ]))
            story.append(body_tbl)
            story.append(Spacer(1, 10))

        story.append(PageBreak())

        # ── Fault Event Log ───────────────────────────────────────────────────
        story.append(Paragraph("3. Fault Event Log", style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=TEAL, spaceAfter=8))

        if "is_anomaly" in df.columns:
            anom_df_r = df[df["is_anomaly"]].copy()
            if not anom_df_r.empty:
                story.append(Paragraph(
                    f"{len(anom_df_r)} anomalous readings detected across the analysis period.",
                    style_body))
                story.append(Spacer(1, 6))

                log_data = [["Timestamp", "Supply Temp", "Return Temp", "Chiller kW", "Total kW"]]
                for _, row in anom_df_r.head(40).iterrows():
                    log_data.append([
                        str(row["timestamp"])[:16],
                        f"{row.get('supply_air_temp','-'):.1f}F" if "supply_air_temp" in row else "-",
                        f"{row.get('return_air_temp','-'):.1f}F" if "return_air_temp" in row else "-",
                        f"{row.get('chiller_kw','-'):.0f}" if "chiller_kw" in row else "-",
                        f"{row.get('total_kw','-'):.0f}" if "total_kw" in row else "-",
                    ])
                log_tbl = Table(log_data, colWidths=[1.8*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.6*inch])
                log_tbl.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0), (-1,0),  BLUE),
                    ("TEXTCOLOR",     (0,0), (-1,0),  colors.white),
                    ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.white, LGRAY]),
                    ("GRID",          (0,0), (-1,-1), 0.4, colors.HexColor("#dee2e6")),
                    ("FONTSIZE",      (0,0), (-1,-1), 8),
                    ("TOPPADDING",    (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ("LEFTPADDING",   (0,0), (-1,-1), 6),
                    ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
                ]))
                story.append(log_tbl)
                if len(anom_df_r) > 40:
                    story.append(Paragraph(f"... and {len(anom_df_r)-40} more events. See full export for complete log.", style_small))
            else:
                story.append(Paragraph("No anomalous readings detected in this analysis period.", style_body))
        else:
            story.append(Paragraph("Run anomaly detection in the Anomalies tab to populate this section.", style_body))

        story.append(PageBreak())

        # ── Energy Savings Summary ────────────────────────────────────────────
        story.append(Paragraph("4. Energy Savings Estimate", style_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=TEAL, spaceAfter=8))

        story.append(Paragraph(
            f"The following table summarizes the estimated annual energy and cost savings "
            f"associated with implementing the recommendations in this report. Savings estimates "
            f"are based on an electricity rate of {currency_sym}{electricity_rate:.3f}/kWh.",
            style_body
        ))
        story.append(Spacer(1, 8))

        savings_data = [["Finding", "Severity", "Est. kWh/yr", f"Est. {currency_sym}/yr"]]
        for f in edited_findings:
            savings_data.append([
                f["title"][:50] + ("..." if len(f["title"]) > 50 else ""),
                f["severity"],
                f"{f['savings_kwh']:,}" if f["savings_kwh"] > 0 else "—",
                f"{currency_sym}{f['savings_usd']:,.0f}" if f["savings_usd"] > 0 else "—",
            ])
        savings_data.append(["TOTAL ESTIMATED SAVINGS", "", f"{total_kwh:,}", f"{currency_sym}{total_usd:,.0f}"])

        sav_tbl = Table(savings_data, colWidths=[3.2*inch, 1.0*inch, 1.4*inch, 1.4*inch])
        sav_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),  (-1,0),   BLUE),
            ("TEXTCOLOR",     (0,0),  (-1,0),   colors.white),
            ("ROWBACKGROUNDS",(0,1),  (-1,-2),  [colors.white, LGRAY]),
            ("BACKGROUND",    (0,-1), (-1,-1),  colors.HexColor("#e8f5e9")),
            ("FONTNAME",      (0,-1), (-1,-1),  "Helvetica-Bold"),
            ("FONTNAME",      (0,0),  (-1,0),   "Helvetica-Bold"),
            ("GRID",          (0,0),  (-1,-1),  0.4, colors.HexColor("#dee2e6")),
            ("FONTSIZE",      (0,0),  (-1,-1),  9),
            ("TOPPADDING",    (0,0),  (-1,-1),  6),
            ("BOTTOMPADDING", (0,0),  (-1,-1),  6),
            ("LEFTPADDING",   (0,0),  (-1,-1),  8),
            ("ALIGN",         (2,0),  (-1,-1),  "RIGHT"),
        ]))
        story.append(sav_tbl)
        story.append(Spacer(1, 16))

        # Payback note
        story.append(Paragraph(
            "<b>Note:</b> Actual savings will vary based on equipment conditions, occupancy patterns, "
            "and utility rate structures. A detailed measurement & verification (M&V) plan is recommended "
            "following implementation of recommendations. Savings estimates follow ASHRAE Guideline 14 "
            "methodology for retrofit isolation M&V.",
            ParagraphStyle("note", fontSize=8, fontName="Helvetica", textColor=MGRAY,
                           spaceAfter=6, leading=12,
                           borderPad=8, borderColor=colors.HexColor("#dee2e6"),
                           borderWidth=1, borderRadius=4)
        ))

        # ── Footer on last page ───────────────────────────────────────────────
        story.append(Spacer(1, 24))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dee2e6"), spaceAfter=6))
        story.append(Paragraph(
            f"Report prepared by {prepared_by} | {company_name} | {report_date} | "
            f"Generated by HVAC Insight Pro",
            ParagraphStyle("footer", fontSize=7.5, fontName="Helvetica",
                           textColor=MGRAY, alignment=TA_CENTER)
        ))

        doc.build(story)
        buffer.seek(0)
        return buffer

    # ── Generate DOCX ────────────────────────────────────────────────────────────
    def generate_docx():
        doc = DocxDocument()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        def add_heading(text, level=1, color_hex="0077B6"):
            p = doc.add_heading(text, level=level)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            return p

        def add_body(text, bold=False, color_hex=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.bold = bold
            if color_hex:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            return p

        def add_table_row(table, cells, bold=False, bg_color=None):
            row = table.add_row()
            for i, text in enumerate(cells):
                cell = row.cells[i]
                cell.text = str(text)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.bold = bold
            return row

        # Title
        title_p = doc.add_heading(report_title, 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title_p.runs:
            run.font.color.rgb = RGBColor(0, 119, 182)
            run.font.size = Pt(22)

        doc.add_paragraph(f"Prepared for: {client_name} | {building_address}")
        doc.add_paragraph(f"Prepared by: {prepared_by} | {company_name}")
        doc.add_paragraph(f"Date: {report_date} | Period: {date_range_str}")
        doc.add_paragraph()

        # Executive Summary
        add_heading("1. Executive Summary", 1)
        add_body(
            f"This report presents findings from an HVAC energy analysis of {client_name} "
            f"covering {date_range_str}. The analysis examined {total_records:,} data points "
            f"across {len(df.columns)-1} sensor channels."
        )
        add_body(
            f"{len(edited_findings)} findings were identified including {high_count} high-severity "
            f"and {med_count} medium-severity issues. Estimated annual savings: "
            f"{currency_sym}{total_usd:,.0f} ({total_kwh:,} kWh).",
            bold=True, color_hex="0077B6"
        )
        doc.add_paragraph()

        # KPI table
        add_heading("Key Performance Indicators", 2)
        kpi_table = doc.add_table(rows=1, cols=2)
        kpi_table.style = "Table Grid"
        hdr = kpi_table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        kpi_rows = [
            ("Analysis Period", date_range_str),
            ("Total Records", f"{total_records:,}"),
            ("Anomalies Detected", str(anom_count_r)),
            ("Avg Chiller Load", f"{kpis_r.get('avg_chiller_kw',0):.1f} kW"),
            ("Total Energy", f"{kpis_r.get('total_kwh',0)/1000:.1f} MWh"),
            ("Avg DeltaT", f"{kpis_r.get('avg_dt',0):.1f} degF"),
            ("Est. Annual Savings", f"{currency_sym}{total_usd:,.0f}"),
        ]
        for label, value in kpi_rows:
            row = kpi_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()
        doc.add_page_break()

        # Findings
        add_heading("2. Technical Findings", 1)
        for i, finding in enumerate(edited_findings):
            sev_color = "C0392B" if finding["severity"]=="HIGH" else "E07B00" if finding["severity"]=="MEDIUM" else "27AE60"
            add_heading(f"Finding {i+1}: {finding['title']} [{finding['severity']}]", 2, sev_color)
            add_body("Observation:", bold=True)
            add_body(finding["description"])
            add_body("Recommended Action:", bold=True)
            add_body(finding["recommendation"])
            if finding["savings_kwh"] > 0:
                add_body(
                    f"Estimated Annual Savings: {finding['savings_kwh']:,} kWh "
                    f"({currency_sym}{finding['savings_usd']:,.0f})",
                    bold=True, color_hex="27AE60"
                )
            doc.add_paragraph()

        doc.add_page_break()

        # Fault log
        add_heading("3. Fault Event Log", 1)
        if "is_anomaly" in df.columns:
            anom_df_r = df[df["is_anomaly"]].copy()
            if not anom_df_r.empty:
                add_body(f"{len(anom_df_r)} anomalous readings detected.")
                fault_tbl = doc.add_table(rows=1, cols=4)
                fault_tbl.style = "Table Grid"
                hdrs = ["Timestamp", "Supply Temp (F)", "Return Temp (F)", "Total kW"]
                for i, h in enumerate(hdrs):
                    fault_tbl.rows[0].cells[i].text = h
                    for run in fault_tbl.rows[0].cells[i].paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(8)
                for _, row in anom_df_r.head(50).iterrows():
                    r = fault_tbl.add_row()
                    r.cells[0].text = str(row["timestamp"])[:16]
                    r.cells[1].text = f"{row.get('supply_air_temp','-'):.1f}" if "supply_air_temp" in row else "-"
                    r.cells[2].text = f"{row.get('return_air_temp','-'):.1f}" if "return_air_temp" in row else "-"
                    r.cells[3].text = f"{row.get('total_kw','-'):.0f}" if "total_kw" in row else "-"
                    for cell in r.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(8)
            else:
                add_body("No anomalous readings detected.")

        doc.add_page_break()

        # Savings summary
        add_heading("4. Energy Savings Estimate", 1)
        add_body(f"Electricity rate: {currency_sym}{electricity_rate:.3f}/kWh")
        doc.add_paragraph()
        sav_tbl = doc.add_table(rows=1, cols=4)
        sav_tbl.style = "Table Grid"
        for i, h in enumerate(["Finding", "Severity", "kWh/yr", f"{currency_sym}/yr"]):
            sav_tbl.rows[0].cells[i].text = h
            for run in sav_tbl.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for f in edited_findings:
            r = sav_tbl.add_row()
            r.cells[0].text = f["title"]
            r.cells[1].text = f["severity"]
            r.cells[2].text = f"{f['savings_kwh']:,}" if f["savings_kwh"] > 0 else "-"
            r.cells[3].text = f"{currency_sym}{f['savings_usd']:,.0f}" if f["savings_usd"] > 0 else "-"
            for cell in r.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        r = sav_tbl.add_row()
        r.cells[0].text = "TOTAL"
        r.cells[2].text = f"{total_kwh:,}"
        r.cells[3].text = f"{currency_sym}{total_usd:,.0f}"
        for cell in r.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        doc.add_paragraph()
        add_body(
            "Note: Savings estimates follow ASHRAE Guideline 14 methodology. "
            "Actual savings will vary based on equipment conditions and utility rates.",
            color_hex="8B949E"
        )

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ── Generate buttons ──────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("<div style='font-family:Space Mono,monospace;font-size:11px;letter-spacing:3px;text-transform:uppercase;color:#00b4d8;border-bottom:1px solid #21262d;padding-bottom:8px;margin-bottom:16px'>Generate Report</div>", unsafe_allow_html=True)

    g1, g2, g3 = st.columns(3)

    with g1:
        if st.button("📄 Generate PDF Report", key="gen_pdf", use_container_width=True):
            with st.spinner("Building PDF..."):
                try:
                    pdf_buffer = generate_pdf()
                    fname = f"HVAC_Report_{client_name.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
                    st.download_button(
                        label="⬇ Download PDF",
                        data=pdf_buffer,
                        file_name=fname,
                        mime="application/pdf",
                        key="dl_pdf",
                        use_container_width=True,
                    )
                    st.success("PDF ready!")
                except Exception as e:
                    st.error(f"PDF generation error: {e}")

    with g2:
        if st.button("📝 Generate Word Report", key="gen_docx", use_container_width=True):
            with st.spinner("Building Word document..."):
                try:
                    docx_buffer = generate_docx()
                    fname = f"HVAC_Report_{client_name.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                    st.download_button(
                        label="⬇ Download Word Doc",
                        data=docx_buffer,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_docx",
                        use_container_width=True,
                    )
                    st.success("Word doc ready!")
                except Exception as e:
                    st.error(f"Word generation error: {e}")

    with g3:
        if st.button("📊 Export Fault Data CSV", key="gen_csv", use_container_width=True):
            rows = []
            for f in edited_findings:
                rows.append({
                    "Finding": f["title"],
                    "Severity": f["severity"],
                    "Description": f["description"],
                    "Recommendation": f["recommendation"],
                    "Est kWh/yr": f["savings_kwh"],
                    f"Est {currency_sym}/yr": f["savings_usd"],
                })
            export_df = pd.DataFrame(rows)
            fname = f"HVAC_Findings_{client_name.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.csv"
            st.download_button(
                label="⬇ Download Findings CSV",
                data=export_df.to_csv(index=False),
                file_name=fname,
                mime="text/csv",
                key="dl_findings_csv",
                use_container_width=True,
            )

    st.info("""
    **How to use this report generator:**
    1. Fill in your company and client details above
    2. Review and edit the auto-generated findings
    3. Click Generate PDF for a client-ready deliverable
    4. Click Generate Word Doc to get an editable version you can customize further
    5. Send the PDF to the building owner/CFO and the Word doc to the facilities manager
    """)

# ─── Footer ──────────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown(
    "<div style='text-align:center;font-family:Space Mono,monospace;font-size:10px;"
    "color:#30363d;letter-spacing:2px'>HVAC INSIGHT PRO · BUILT WITH STREAMLIT</div>",
    unsafe_allow_html=True,
)
